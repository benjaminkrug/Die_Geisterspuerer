"""
Build a properly formatted DOCX for KDP Taschenbuch (paperback, 5x8 inches)
from Manuskript_Band1_Komplett.md.
Handles: page size, mirror margins, page numbers, chapter headings,
paragraphs with first-line indent, scene breaks, italic/bold text,
front matter, back matter.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Benjamin Krug"
INPUT_FILE = "Manuskript_Band1_Komplett.md"
OUTPUT_FILE = "KDP_Band1_Manuskript.docx"

SCENE_BREAK_SYMBOL = "\u2726  \u2726  \u2726"


def setup_page(section):
    """Set 5x8 inch page size and KDP margins for a section."""
    section.page_width = Inches(5)
    section.page_height = Inches(8)
    section.left_margin = Inches(0.75)    # inner margin (with mirror)
    section.right_margin = Inches(0.5)    # outer margin (with mirror)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.625)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.3)


def setup_mirror_margins(doc):
    """Enable mirror margins for bookbinding (inner/outer swap on even/odd pages)."""
    settings_element = doc.settings.element
    mirror = OxmlElement('w:mirrorMargins')
    settings_element.append(mirror)


def add_page_number_footer(section):
    """Add centered page number to section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PAGE field: begin
    run1 = p.add_run()
    run1.font.name = "Georgia"
    run1.font.size = Pt(10)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fld_begin)

    # PAGE field: instruction
    run2 = p.add_run()
    run2.font.name = "Georgia"
    run2.font.size = Pt(10)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._r.append(instr)

    # PAGE field: end
    run3 = p.add_run()
    run3.font.name = "Georgia"
    run3.font.size = Pt(10)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld_end)


def create_styles(doc):
    """Set up document styles for Taschenbuch formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.5)  # ~1.2em at 11pt
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_lines(doc, count=1):
    """Add empty paragraphs."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    """Add a paragraph with inline formatting (*italic*, **bold**)."""
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)

    # Parse inline markdown formatting
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.findall(pattern, text)

    for full_match, bold_text, italic_text, plain_text in parts:
        if bold_text:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif italic_text:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif plain_text:
            run = p.add_run(plain_text)
            run.font.name = "Georgia"
            run.font.size = Pt(11)

    return p


def add_front_matter(doc):
    """Add title page, half-title, and copyright page."""
    # --- Half title ---
    add_blank_lines(doc, 6)
    add_centered_text(doc, "Die Geisterspürer", font_size=18, bold=True)
    add_page_break(doc)

    # --- Full title page ---
    add_blank_lines(doc, 3)
    add_centered_text(doc, "Die Geisterspürer", font_size=22, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Das Haus, das flüstert", font_size=16, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Band 1", font_size=12)
    add_blank_lines(doc, 2)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # --- Copyright page ---
    add_blank_lines(doc, 8)
    add_centered_text(doc, "Die Geisterspürer – Das Haus, das flüstert", font_size=10, bold=True)
    add_centered_text(doc, "Band 1", font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = Pt(9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)


def add_back_matter(doc):
    """Add back matter: review request, Band 2 teaser, series list, cross-references."""
    add_page_break(doc)

    # --- 1. Rezensions-Bitte ---
    add_blank_lines(doc, 3)
    add_centered_text(doc, "\u201eDas Haus, das fl\u00fcstert\u201c hat dir gefallen?", font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # --- 2. Teaser Band 2 ---
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Weiterlesen? Hier kommt eine Vorschau auf Band 2:", font_size=11, italic=True)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Geisterspürer", font_size=18, bold=True)
    add_centered_text(doc, "Der Friedhof ohne Namen", font_size=14, italic=True)
    add_centered_text(doc, "Band 2", font_size=11)
    add_blank_lines(doc, 2)

    Q = "\u201e"
    E = "\u201c"
    teaser_paragraphs = [
        "Schatten wollte nicht auf den Friedhof.",
        f"Das war neu. Bei Lina hatte er sie nach oben gezerrt, in die verlassene Wohnung. Er hatte gekratzt und geknurrt und nicht aufgeh\u00f6rt, bis Nora ihm gefolgt war. Aber vor dem rostigen Eisentor des alten Friedhofs stand er still. Steif. Sein Fell ges\u00e4ubt, sein Blick starr \u2014 und zum ersten Mal wollte er zur\u00fcck.",
        f"{Q}Er hat Angst{E}, fl\u00fcsterte Theo.",
        f"Nora schluckte. Sie hatte Schatten knurren gesehen. Bellen. Sie hatte ihn vor sich und Theo stellen sehen wie einen Schild. Aber Angst? Noch nie.",
        f"{Q}Wir k\u00f6nnen umdrehen{E}, sagte Theo. Er meinte es ernst. Kein Witz diesmal.",
        f"Nora schaute auf die Karte in ihrer Hand. Frau Silbers Handschrift. *Heinrich M. \u2014 1847. W\u00fctend. Grab.* Hundertachtzig Jahre. Allein zwischen den Grabsteinen. W\u00e4hrend die Stadt um ihn herum lebte und lachte und ihn vergessen hatte.",
        f"{Q}Nein{E}, sagte sie. {Q}Wir gehen rein.{E}",
        f"{Q}Das sagt das M\u00e4dchen, das vor drei Wochen nicht mal an Geister geglaubt hat.{E}",
        f"{Q}Vor drei Wochen wusste ich auch nicht, dass unser Hund \u00fcbernat\u00fcrlich ist.{E}",
        "Das Tor quietschte. Der Klang hallte zwischen den Grabsteinen.",
        "Nora ging voraus. Kies unter ihren Schuhen. Moos an den Steinen. Stille, die sich anf\u00fchlte wie Watte in den Ohren.",
        "Schatten folgte. Langsam. Jeder Schritt ein Kampf gegen seinen Instinkt.",
        f"Dann blieb er stehen. Am Ende der dritten Reihe. Vor einem Grabstein, der anders war als alle anderen. Gr\u00f6\u00dfer. Dunkler. Schwarzer Stein, glatt poliert \u2014 als w\u00fcrde ihn jemand pflegen.",
        "Aber es stand kein Name darauf. Kein Datum. Nichts.",
        f"{Q}Ein Grabstein ohne Namen{E}, murmelte Theo. {Q}Wer liegt da?{E}",
        "Nora kniete sich hin. Vor dem Stein lag etwas im Gras. Klein. Metallisch. Ein Schl\u00fcssel. Alt, verrostet, aber sorgf\u00e4ltig hingelegt. Als h\u00e4tte ihn jemand erst gestern dort platziert.",
        f"Und das Gras vor dem Grabstein war niedergedr\u00fcckt. In einem Kreis. Genau so gro\u00df wie ein Mensch, der steht.",
        f"{Q}Jemand steht hier{E}, fl\u00fcsterte sie. {Q}Jede Nacht. An genau dieser Stelle.{E}",
        f"{Q}Seit hundertachtzig Jahren{E}, sagte Theo. Seine Stimme war kaum zu h\u00f6ren.",
        f"Schatten hob den Kopf. Sein Knurren kam so tief aus seiner Brust, dass Nora es in den F\u00fc\u00dfen sp\u00fcrte. Sein ganzer K\u00f6rper vibrierte.",
        f"Dann tat er etwas, das er noch nie getan hatte. Er wich zur\u00fcck. Schritt f\u00fcr Schritt. Ohne den Blick von dem Grabstein zu nehmen.",
        f"Die Temperatur fiel. Noras Atem wurde sichtbar. Wei\u00dfe W\u00f6lkchen vor ihrem Mund \u2014 mitten im Juli.",
        f"Und auf dem schwarzen Grabstein erschienen Buchstaben. Langsam. Einer nach dem anderen. Wie von einer unsichtbaren Hand in den Stein geritzt.",
        "G E H T.",
        f"Theo griff nach Noras Arm. Seine Finger waren eiskalt.",
        f"Die Buchstaben verschwanden. Neue kamen.",
        f"O D E R.",
        f"Stille. Schattens Knurren erstarb. Die Luft stand still. Sogar der Wind h\u00f6rte auf.",
        f"Dann das letzte Wort. Gr\u00f6\u00dfer als die anderen. Tiefer in den Stein gedr\u00fcckt.",
        "B L E I B T.",
        f"Nora sp\u00fcrte ihr Herz bis in die Fingerspitzen. Hinter ihr wich Theo zur\u00fcck. Vor ihr leuchteten die Buchstaben im schwarzen Stein. Und neben ihr knurrte Schatten \u2014 nicht den Grabstein an.",
        "Sondern etwas hinter ihnen.",
    ]

    for i, para in enumerate(teaser_paragraphs):
        add_formatted_paragraph(doc, para, first_line_indent=(i > 0))

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Erscheint 2026 auf Amazon.", font_size=11, italic=True)

    # --- 3. Serien-Übersicht ---
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Geisterspürer \u2014 Alle Bände", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    for band_num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(
            doc,
            f"**Band {band_num}:** {title}",
            first_line_indent=False,
        )

    # --- 4. Cross-Verweis andere Serien ---
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Mehr spannende Abenteuer von Benjamin Krug:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    # Herrenhaus-Detektive
    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schl\u00fcssel \u2014 und hinter der verschlossenen "
        "T\u00fcr wartet ein Geheimnis, das seit drei\u00dfig Jahren niemand "
        "l\u00fcften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)

    # Chrono-Agenten
    add_centered_text(doc, "Die Chrono-Agenten", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Drei Kinder. Ein Tablet, das die Zeit \u00f6ffnet. Und ein "
        "Countdown, der nicht aufh\u00f6rt zu laufen.",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "Leo, Mila und Ben sind Agenten wider Willen \u2014 geschickt in die "
        "gef\u00e4hrlichsten Momente der Geschichte, um zu retten, was jemand "
        "zerst\u00f6ren will. Ihr Gegner hei\u00dft Codex. Und er kennt ihre Namen.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "F\u00fcr alle, die Geschichte spannend finden \u2014 und wissen wollen, "
        "was passiert, wenn man sie ver\u00e4ndert.",
        first_line_indent=False,
    )


def add_chapter_heading(doc, title):
    """Add a chapter heading with page break and precise spacing."""
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.name = "Georgia"
    run.bold = True


def add_scene_break(doc):
    """Add an ornamental scene break."""
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)


def parse_and_build(doc, content):
    """Parse markdown content and build the DOCX."""
    match = re.search(r'^# Kapitel 1', content, re.MULTILINE)
    if not match:
        raise ValueError("Could not find '# Kapitel 1' in manuscript")
    body = content[match.start():]

    # Remove ENDE BAND 1 and trailing ---
    body = re.sub(r'\n---\n\n\*\*ENDE BAND 1\*\*\n\n---\n*', '', body)
    body = re.sub(r'\n\*\*ENDE BAND 1\*\*\n*', '', body)

    lines = body.split('\n')

    is_first_para_after_heading = False
    is_first_para_after_break = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Chapter heading
        if line.startswith('# Kapitel '):
            title = line[2:]  # Remove "# "
            add_chapter_heading(doc, title)
            is_first_para_after_heading = True
            i += 1
            continue

        # Scene break
        if line == '---':
            add_scene_break(doc)
            is_first_para_after_break = True
            i += 1
            continue

        # Regular paragraph
        no_indent = is_first_para_after_heading or is_first_para_after_break
        add_formatted_paragraph(doc, line, first_line_indent=not no_indent)
        is_first_para_after_heading = False
        is_first_para_after_break = False
        i += 1


def main():
    # Read manuscript
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # --- Section 1: Front matter (no page numbers) ---
    section_front = doc.sections[0]
    setup_page(section_front)

    # Suppress footer on front matter
    header = section_front.header
    header.is_linked_to_previous = False
    footer = section_front.footer
    footer.is_linked_to_previous = False

    # Enable mirror margins for bookbinding
    setup_mirror_margins(doc)

    # Set up styles
    create_styles(doc)

    # Build front matter
    add_front_matter(doc)

    # --- Section 2: Body + back matter (with page numbers) ---
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_body = doc.sections[-1]
    setup_page(section_body)

    # Add page numbers to body section footer
    add_page_number_footer(section_body)

    # Suppress header in body section
    header_body = section_body.header
    header_body.is_linked_to_previous = False

    # Parse and build body content
    parse_and_build(doc, content)

    # Add back matter
    add_back_matter(doc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"{OUTPUT_FILE} erstellt!")

    # Stats
    para_count = len(doc.paragraphs)
    print(f"Absätze: {para_count}")

    chapters = len(re.findall(r'# Kapitel \d+', content))
    print(f"Kapitel: {chapters}")


if __name__ == "__main__":
    main()
