"""
Baut ein KDP-fertiges DOCX + PDF (Taschenbuch 6x9 Zoll) fuer Staffel 2, Band 1
("Der Gast, der blieb") aus Manuskript_S2-1_Komplett.md.

Abgeleitet von build_taschenbuch_docx_band5.py. Typografie unveraendert
uebernommen (deutsche Anfuehrungszeichen paarweise, Kapitaelchen-Auftakt,
Gedankenstriche). PDF-Konvertierung via LibreOffice (soffice) oder Word.

Verwendung:
    python Scripts/build_taschenbuch_docx_s2_1.py            # DOCX + PDF bauen
    python Scripts/build_taschenbuch_docx_s2_1.py --pruefen  # nur pruefen


★ NUR-PRUEFEN-MODUS — und warum es ihn gibt

  --pruefen laeuft OHNE python-docx. Es fuehrt genau die Textaufbereitung aus,
  die auch der echte Build benutzt (ENDE-Marker strippen, Ereignisfolge bauen),
  und meldet, was im gedruckten Buch landen wuerde.

  Das ist kein Komfort-Feature. **Band 5s einziger echter Druckfehler sass genau
  hier**: Kapitel_18.md endete mit dem Marker "**ENDE**", der bis ins fertige
  DOCX durchschlug — direkt vor die Rezensions-Bitte — und erst beim Pruefen des
  fertigen Dokuments auffiel. Alles, was diesen Fehler haette fangen koennen,
  braucht kein Word: es ist Textverarbeitung. Also laeuft es auch ohne.


★ UNTERSCHIEDE ZU BAND 5 — jeder mit Grund:

1. ENDE-REGEX FUER ZWEI MARKER.
   Das S2-1-Manuskript enthaelt BEIDE:
     - "**ENDE**"                        (Schluss von Kapitel 16, im Manuskript)
     - "**ENDE BAND 1 · DIE GEBUNDENEN**" (Build-Marker)
   ‼️ BEIDE gehoeren NICHT in den Druck. In den Kapiteldateien von Band 1 bis 4
   kommt "**ENDE**" ueberhaupt nicht vor; nur Band 5 fuehrt ihn, und dessen
   Taschenbuch-Skript streicht ihn ausdruecklich wieder heraus. Kein gedrucktes
   Geisterspuerer-Buch hat ein alleinstehendes "ENDE".

2. KEIN QR-CODE, SOLANGE ES KEINE ASIN GIBT.
   S2-1 ist nicht veroeffentlicht. Band 5s Skript haette hier stillschweigend
   ein fehlendes Bild uebersprungen und ein Buch gebaut, in dem "Einfach den
   Code scannen" steht und kein Code folgt. Hier: BAND_ASIN = None schaltet den
   ganzen Block ab, und main() sagt es laut.

3. KEIN TEASER auf den naechsten Band.
   S2-2 ist nicht geschrieben. Die Regel steht seit Band 5 fest: Von Band 4s
   12 Leseproben-Absaetzen existieren 0 in Band 5; von Band 3s 16 existiert
   genau 1 in Band 4. Beide waren spekulativ geschrieben. **Erst NACH dem
   naechsten Manuskript, und gegen den echten Text pruefen.**

4. SERIENUEBERSICHT ZEIGT NUR KAUFBARE TITEL.
   Band 5 listete alle fuenf Baende seiner Staffel. Fuer S2-1 waeren S2-2 bis
   S2-5 unkaufbar — die Regel aus Band 5s Skript gilt unveraendert:
   ‼️ Nie einen Titel bewerben, den es noch nicht zu kaufen gibt.
   Stattdessen die fuenf Baende der ersten Staffel als Backlist, mit dem
   ausdruecklichen Hinweis, dass die Reihenfolge egal ist (freier Einstieg,
   PLAN_Staffel2.md Abschnitt 3).

5. TITELSEITE MIT REIHENNAME.
   "Die Geisterspuerer" hat jetzt zwei Baende mit der Nummer 1. Die Titelseite
   traegt deshalb "Die Gebundenen · Band 1" (der Name steht so auch auf dem
   Cover, PLAN_Staffel2.md Abschnitt 10).
   ⚠️ Das Wort "Staffel" taucht NIRGENDS im Buch auf — das ist ein internes
   Planungswort, der Leser kennt es nicht (Regel aus Band 5).

Unveraendert aus Band 5 uebernommen (dort begruendet):
   - defensiver Skip fuer einen Szenentrenner direkt vor einer Kapitelueberschrift
   - Frontmatter wird IMPORTIERT, nicht kopiert (single source of truth)
"""

import os
import re
import sys
import importlib.util

AUTHOR = "Benjamin Krug"
SERIES_TITLE = "Die Geisterspürer"
STAFFEL_TITLE = "Die Gebundenen"
BAND_NUM = 1
BAND_TITLE = "Der Gast, der blieb"
BAND_SUBTITLE = f"{STAFFEL_TITLE} · Band {BAND_NUM}"
ERWARTETE_KAPITEL = 16

# Kein NEXT_BAND — S2-2 ist nicht geschrieben (s. Punkt 3 oben).

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.join(_SCRIPT_DIR, "..")

INPUT_FILE = os.path.join(_ROOT, "Staffel2", "S2-1", "Manuskript",
                          "Manuskript_S2-1_Komplett.md")
OUTPUT_DIR = os.path.join(_ROOT, "Output", "S2-1")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "KDP_S2-1_Manuskript.docx")
OUTPUT_PDF = os.path.join(OUTPUT_DIR, "KDP_S2-1_Manuskript.pdf")

SCENE_BREAK_SYMBOL = "✦  ✦  ✦"

# --- Amazon-Rezension ---------------------------------------------------------
# ⚠️ S2-1 ist noch nicht veroeffentlicht, es gibt keine ASIN. Solange das so ist,
#    baut das Skript die Rezensions-Seite OHNE QR-Block (s. Punkt 2 oben).
#    Nach der Veroeffentlichung: ASIN eintragen, Scripts/build_qr_rezension.py
#    laufen lassen, Pfad unten setzen, neu bauen.
BAND_ASIN = None
REVIEW_URL = f"https://www.amazon.de/review/create-review?asin={BAND_ASIN}" if BAND_ASIN else None
QR_IMAGE = os.path.join(_ROOT, "Staffel2", "S2-1", "Cover", "qr_rezension_s2_1.png")

Q = "„"
E = "“"

SMALLCAPS_WORDS = 4


# ══════════════════════════════════════════════════════════════════════════════
# TEIL 1 — TEXTAUFBEREITUNG (ohne python-docx, deshalb auch --pruefen-tauglich)
# ══════════════════════════════════════════════════════════════════════════════

# Beide Marker, in dieser Reihenfolge. Die \n+-Form ist Absicht: eine exakte
# \n\n-Form greift nicht mehr, sobald der Build die Zahl der Leerzeilen aendert.
ENDE_BUILD_MARKER = (
    r'\n+---\n+\*\*ENDE BAND ' + str(BAND_NUM) + r' · ' +
    STAFFEL_TITLE.upper() + r'\*\*\n+---\n*'
)
ENDE_KAPITEL_MARKER = r'(?m)^\s*\*\*ENDE\*\*\s*$\n?'


def strip_ende_marker(body: str) -> str:
    body = re.sub(ENDE_BUILD_MARKER, '', body)
    body = re.sub(ENDE_KAPITEL_MARKER, '', body)
    return body


def uebrige_ende_zeilen(body: str) -> list:
    """Band-5-Guard: Was nach dem Strippen noch wie ein ENDE-Marker aussieht,
    wuerde im Buch landen. Nur die exakte Marker-Zeile, niemals das Wort 'Ende'
    im Fliesstext."""
    return [z for z in body.split('\n')
            if re.fullmatch(r'\s*\*{0,2}ENDE[^*]*\*{0,2}\s*', z)]


def bereite_body_vor(content: str) -> str:
    m = re.search(r'^# Kapitel 1', content, re.MULTILINE)
    if not m:
        raise ValueError("Konnte '# Kapitel 1' nicht finden.")
    body = strip_ende_marker(content[m.start():])

    rest = uebrige_ende_zeilen(body)
    if rest:
        raise ValueError(
            f"ENDE-Marker nicht sauber entfernt, wuerde im Buch landen: {rest!r} "
            f"(Unterschied 1 im Dateikopf — Regex pruefen.)"
        )
    return body


def ereignisfolge(body: str) -> list:
    """Dieselbe Reihenfolge, die parse_and_build spaeter setzt — nur ohne Word.
    Erlaubt es, das gedruckte Ergebnis zu pruefen, bevor es existiert."""
    lines = [z.strip() for z in body.split('\n')]
    ereignisse = []
    for i, z in enumerate(lines):
        if not z:
            continue
        if z.startswith('# Kapitel '):
            ereignisse.append(('KAPITEL', z))
        elif z == '---':
            # LANDMINE 1 (aus Band 5): Trenner direkt vor einer Kapitelueber-
            # schrift wird uebersprungen — er stuende hinter dem Cliffhanger.
            j = i + 1
            while j < len(lines) and not lines[j]:
                j += 1
            if j < len(lines) and lines[j].startswith('# Kapitel '):
                continue
            ereignisse.append(('TRENNER', z))
        else:
            ereignisse.append(('TEXT', z))
    return ereignisse


# ── Typografie ────────────────────────────────────────────────────────────────

def typo_quotes(text: str) -> str:
    """Gerade ASCII-Anfuehrungszeichen zu deutschen „…“, paarweise."""
    out = []
    inside = False
    for idx, ch in enumerate(text):
        if ch == '"':
            out.append('“' if inside else '„')
            inside = not inside
        elif ch == "'":
            # Apostroph nur innerhalb oder am Wortende ersetzen („sag's", „hab's").
            # Ein fuehrendes ' bleibt stehen — es waere ein Anfuehrungszeichen.
            prev = text[idx - 1] if idx > 0 else ' '
            out.append('’' if prev.isalpha() else ch)
        else:
            out.append(ch)
    return "".join(out)


def typo_dashes(text: str) -> str:
    text = re.sub(r'\s—\s', ' – ', text)
    text = re.sub(r'\s—', ' –', text)
    text = re.sub(r'—\s', '– ', text)
    return text


def apply_typography(text: str) -> str:
    return typo_dashes(typo_quotes(text))


# ── Frontmatter — IMPORTIERT, nicht kopiert (Band-5-Lehre) ────────────────────
# Single source of truth: build_manuskript_komplett_s2_1.py. Band 1-4 definierten
# Widmung und Epigraph in ZWEI Dateien; laufen die auseinander, weicht das
# gedruckte Buch vom Manuskript ab, ohne dass es jemand merkt.
def _lade_frontmatter():
    pfad = os.path.join(_SCRIPT_DIR, "build_manuskript_komplett_s2_1.py")
    spec = importlib.util.spec_from_file_location("_s21_kompilat", pfad)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)          # main() ist durch __main__-Guard geschuetzt
    return mod.WIDMUNG_ZEILEN, mod.EPIGRAPH_ZEILEN, mod.EPIGRAPH_SOURCE


WIDMUNG_ZEILEN, EPIGRAPH_ZEILEN, EPIGRAPH_SOURCE = _lade_frontmatter()


# ══════════════════════════════════════════════════════════════════════════════
# TEIL 2 — DOCX (braucht python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def _docx():
    """Laedt python-docx erst, wenn wirklich gebaut wird. So laeuft --pruefen
    auch dort, wo die Bibliothek fehlt."""
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    return dict(Document=Document, Pt=Pt, Cm=Cm, Inches=Inches,
                AL=WD_ALIGN_PARAGRAPH, BR=WD_BREAK, SEC=WD_SECTION_START,
                qn=qn, Ox=OxmlElement)


D = None   # wird in main() gefuellt


def setup_page(section):
    section.page_width = D['Inches'](6)
    section.page_height = D['Inches'](9)
    section.left_margin = D['Inches'](0.875)   # innen (Bindung)
    section.right_margin = D['Inches'](0.625)  # aussen
    section.top_margin = D['Inches'](0.75)
    section.bottom_margin = D['Inches'](0.75)
    section.header_distance = D['Inches'](0)
    section.footer_distance = D['Inches'](0.35)


def setup_mirror_margins(doc):
    doc.settings.element.append(D['Ox']('w:mirrorMargins'))


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = D['AL'].CENTER
    for fld_type, content in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        run = p.add_run()
        run.font.name = "Georgia"
        run.font.size = D['Pt'](10)
        if fld_type == 'instrText':
            instr = D['Ox']('w:instrText')
            instr.set(D['qn']('xml:space'), 'preserve')
            instr.text = content
            run._r.append(instr)
        else:
            fld = D['Ox']('w:fldChar')
            fld.set(D['qn']('w:fldCharType'), fld_type)
            run._r.append(fld)


def create_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = D['Pt'](11)
    pf = style.paragraph_format
    pf.space_before = D['Pt'](0)
    pf.space_after = D['Pt'](0)
    pf.line_spacing = 1.5
    pf.first_line_indent = D['Cm'](0.6)
    pf.alignment = D['AL'].JUSTIFY
    pf.widow_control = True


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = D['Cm'](0)
    p.add_run().add_break(D['BR'].PAGE)


def add_blank_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = D['Cm'](0)
        p.paragraph_format.space_before = D['Pt'](0)
        p.paragraph_format.space_after = D['Pt'](0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = D['AL'].CENTER
    p.paragraph_format.first_line_indent = D['Cm'](0)
    p.paragraph_format.space_before = D['Pt'](0)
    p.paragraph_format.space_after = D['Pt'](0)
    run = p.add_run(text)
    run.font.size = D['Pt'](font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_centered_image(doc, image_path, width_inches=1.6):
    if not os.path.isfile(image_path):
        print(f"  WARNUNG: QR-Bild nicht gefunden, wird uebersprungen: {image_path}")
        return None
    p = doc.add_paragraph()
    p.alignment = D['AL'].CENTER
    p.paragraph_format.first_line_indent = D['Cm'](0)
    p.add_run().add_picture(image_path, width=D['Inches'](width_inches))
    return p


def _add_runs_with_markdown(p, text):
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for _full, bold_text, italic_text, plain_text in re.findall(pattern, text):
        if bold_text:
            run = p.add_run(bold_text); run.bold = True
        elif italic_text:
            run = p.add_run(italic_text); run.italic = True
        elif plain_text:
            run = p.add_run(plain_text)
        else:
            continue
        run.font.name = "Georgia"
        run.font.size = D['Pt'](11)


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    text = apply_typography(text)
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = D['Cm'](0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = D['Cm'](0)
    _add_runs_with_markdown(p, text)
    return p


# ── Front matter ──────────────────────────────────────────────────────────────

def add_front_matter(doc):
    # Halbtitel
    add_blank_lines(doc, 8)
    add_centered_text(doc, SERIES_TITLE, font_size=20, bold=True)
    add_page_break(doc)

    # Titelseite
    add_blank_lines(doc, 4)
    add_centered_text(doc, SERIES_TITLE, font_size=24, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, STAFFEL_TITLE, font_size=13)
    add_blank_lines(doc, 1)
    add_centered_text(doc, BAND_TITLE, font_size=17, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Band {BAND_NUM}", font_size=12)
    add_blank_lines(doc, 3)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # Impressum
    add_blank_lines(doc, 10)
    add_centered_text(doc, f"{SERIES_TITLE} – {BAND_TITLE}", font_size=10, bold=True)
    add_centered_text(doc, BAND_SUBTITLE, font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = D['Pt'](9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)

    # Widmung (eigene Seite)
    add_page_break(doc)
    add_blank_lines(doc, 12)
    for zeile in WIDMUNG_ZEILEN:
        add_centered_text(doc, zeile, font_size=12, italic=True)
        add_blank_lines(doc, 1)

    # Epigraph (eigene Seite)
    add_page_break(doc)
    add_blank_lines(doc, 12)
    for zeile in EPIGRAPH_ZEILEN:
        add_centered_text(doc, zeile, font_size=12, italic=True)
        add_blank_lines(doc, 1)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"— {EPIGRAPH_SOURCE}", font_size=10, italic=True)


# ── Back matter ───────────────────────────────────────────────────────────────

def add_back_matter(doc):
    add_page_break(doc)

    # 1. Rezensions-Bitte
    add_blank_lines(doc, 3)
    add_centered_text(doc, f"{Q}{BAND_TITLE}{E} hat dir gefallen?", font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 2)

    # QR nur, wenn es eine ASIN GIBT (Unterschied 2 im Dateikopf).
    # Sonst stuende "Einfach den Code scannen" ohne Code im Buch.
    if BAND_ASIN:
        add_centered_text(doc, "Einfach den Code scannen und eine Bewertung dalassen:",
                          font_size=11, italic=True)
        add_blank_lines(doc, 1)
        add_centered_image(doc, QR_IMAGE, width_inches=1.6)
        add_blank_lines(doc, 1)
        add_centered_text(doc, "(Handykamera auf den Code halten – der Link öffnet sich von selbst.)",
                          font_size=9, italic=True)
        add_blank_lines(doc, 2)

    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # 2. KEIN TEASER — S2-2 ist nicht geschrieben (Unterschied 3 im Dateikopf).

    # 3. Serienuebersicht — nur kaufbare Titel (Unterschied 4 im Dateikopf)
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, f"{SERIES_TITLE} — die ersten fünf Bände",
                      font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Nora, Theo und Schatten hatten schon einmal zu tun. Fünf Fälle, "
        "bevor dieser hier anfing.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    for num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(doc, f"**Band {num}:** {title}", first_line_indent=False)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Man kann sie in jeder Reihenfolge lesen — auch nach diesem Buch.",
                      font_size=10, italic=True)

    # 4. Cross-Verweis (alle beworbenen Titel sind veroeffentlicht)
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Mehr spannende Abenteuer von {AUTHOR}:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    add_centered_text(doc, "Die Geisterspürer – als Spielbuch", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Du kennst die Geschichte. Aber was wäre passiert, wenn Nora damals "
        "nicht in den Keller gegangen wäre?",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        f"In {Q}Das Haus, das flüstert – Grusel-Spielbuch{E} entscheidest du an jedem "
        "Wendepunkt selbst. Folgst du dem Hund die Treppe hinauf? Gehst du nachts "
        "in den Keller? Liest du das Tagebuch allein — oder mit Theo?",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "24 verschiedene Enden. Und ein geheimes, das nur die Mutigsten finden.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Grusel-Spielbuch ab 10 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schlüssel — und hinter der verschlossenen Tür wartet ein "
        "Geheimnis, das seit dreißig Jahren niemand lüften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Herrenhaus-Detektive – als Spielbuch", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        f"Zwei Fälle, in denen du selbst ermittelst: {Q}Das verbotene Herrenhaus{E} "
        f"und {Q}Das Geheimnis des Brunnens{E}.",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "Welcher Spur folgst du zuerst? Wem glaubst du? Jede Entscheidung führt "
        "dich auf einem anderen Weg durch den Fall — und zu einem anderen Ende.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Detektiv-Spielbücher ab 8 Jahren.", font_size=11, italic=True)


# ── Kapitel ───────────────────────────────────────────────────────────────────

def add_chapter_heading(doc, title):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = D['AL'].CENTER
    p.paragraph_format.first_line_indent = D['Cm'](0)
    p.paragraph_format.space_before = D['Pt'](60)
    p.paragraph_format.space_after = D['Pt'](36)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.font.size = D['Pt'](14)
    run.font.name = "Georgia"
    run.bold = True


def add_scene_break(doc):
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)


def add_chapter_first_paragraph(doc, text):
    """Erste SMALLCAPS_WORDS Woerter in Kapitaelchen. Schutzklausel: Beginnt der
    Absatz mit kursivem Markdown, bleibt er kursiv."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = D['Cm'](0)

    if text.lstrip().startswith('*'):
        _add_runs_with_markdown(p, text)
        return p

    m = re.match(r'^(\s*)(.*)$', text)
    leading_ws, rest = m.group(1), m.group(2)
    tokens = rest.split(' ')
    head = ' '.join(tokens[:SMALLCAPS_WORDS])
    tail_tokens = tokens[SMALLCAPS_WORDS:]
    tail = (' ' + ' '.join(tail_tokens)) if tail_tokens else ''

    if '*' in head:
        _add_runs_with_markdown(p, text)
        return p

    if leading_ws:
        run = p.add_run(leading_ws)
        run.font.name = "Georgia"
        run.font.size = D['Pt'](11)

    run = p.add_run(head)
    run.font.name = "Georgia"
    run.font.size = D['Pt'](11)
    run.font.small_caps = True

    if tail:
        _add_runs_with_markdown(p, tail)
    return p


def parse_and_build(doc, content):
    body = bereite_body_vor(content)
    erste_nach_kapitel = False
    erste_nach_trenner = False

    for art, zeile in ereignisfolge(body):
        if art == 'KAPITEL':
            add_chapter_heading(doc, zeile[2:])
            erste_nach_kapitel, erste_nach_trenner = True, False
        elif art == 'TRENNER':
            add_scene_break(doc)
            erste_nach_kapitel, erste_nach_trenner = False, True
        else:
            text = apply_typography(zeile)
            if erste_nach_kapitel:
                add_chapter_first_paragraph(doc, text)
            else:
                add_formatted_paragraph(
                    doc, text,
                    first_line_indent=not (erste_nach_kapitel or erste_nach_trenner))
            erste_nach_kapitel = erste_nach_trenner = False


# ── PDF ───────────────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path, pdf_path):
    import subprocess
    out_dir = os.path.dirname(os.path.abspath(pdf_path))
    for soffice in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
    ):
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir,
                 os.path.abspath(docx_path)],
                check=True, capture_output=True,
            )
            erzeugt = os.path.join(
                out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if os.path.exists(erzeugt):
                if os.path.abspath(erzeugt) != os.path.abspath(pdf_path):
                    os.replace(erzeugt, os.path.abspath(pdf_path))
                print(f"PDF erstellt (LibreOffice): {pdf_path}")
                return True
        except Exception:
            continue
    print(f"PDF-Konvertierung fehlgeschlagen. DOCX ist fertig: {docx_path}")
    return False


# ══════════════════════════════════════════════════════════════════════════════
# PRUEFEN + MAIN
# ══════════════════════════════════════════════════════════════════════════════

def pruefen(content: str) -> int:
    print("=" * 74)
    print("PRUEFUNG DES DRUCKBILDS — ohne python-docx")
    print("=" * 74)
    fehler = []

    body = bereite_body_vor(content)          # wirft bei ENDE-Resten
    print("ENDE-Marker: beide entfernt, keine Restzeile ✓")

    ev = ereignisfolge(body)
    kapitel = [z for a, z in ev if a == 'KAPITEL']
    trenner = [1 for a, _ in ev if a == 'TRENNER']
    print(f"Kapitel im Druck: {len(kapitel)} (erwartet {ERWARTETE_KAPITEL})")
    print(f"Echte Szenentrenner: {len(trenner)}")
    if len(kapitel) != ERWARTETE_KAPITEL:
        fehler.append(f"{len(kapitel)} Kapitel statt {ERWARTETE_KAPITEL}")

    for i in range(len(ev) - 1):
        if ev[i][0] == 'TRENNER' and ev[i + 1][0] == 'KAPITEL':
            fehler.append(f"Fehl-Ornament vor {ev[i+1][1]!r}")
        if ev[i][0] == 'KAPITEL' and ev[i + 1][0] == 'TRENNER':
            fehler.append(f"Ornament direkt nach {ev[i][1]!r}")
    print(f"Fehl-Ornamente: {sum(1 for f in fehler if 'Ornament' in f)}")

    # Letzte Textzeile des Buchs — dort landete bei Band 5 der ENDE-Marker
    letzte = [z for a, z in ev if a == 'TEXT'][-1]
    print(f"Letzte Zeile vor dem Nachwort: {letzte!r}")
    if re.search(r'ENDE', letzte):
        fehler.append("Letzte Zeile enthaelt einen ENDE-Marker")

    # Typografie stichprobenartig
    probe = apply_typography('"Warum reicht das nicht?" — sagte sie.')
    print(f"Typografie-Probe: {probe}")
    if '"' in probe:
        fehler.append("Gerade Anfuehrungszeichen ueberlebt die Typografie")

    print(f"\nFrontmatter (importiert aus build_manuskript_komplett_s2_1.py):")
    print(f"   Widmung : {WIDMUNG_ZEILEN[0]}")
    print(f"   Epigraph: {EPIGRAPH_ZEILEN[0]}")

    if BAND_ASIN is None:
        print("\n⚠️  KEINE ASIN — die Rezensions-Seite wird OHNE QR-Code gebaut.")
        print("    Nach der Veroeffentlichung: BAND_ASIN setzen,")
        print("    Scripts/build_qr_rezension.py laufen lassen, neu bauen.")
    print("⚠️  KEIN TEASER auf S2-2 — der Band ist nicht geschrieben.")

    print()
    if fehler:
        print("!! BEFUNDE:")
        for f in dict.fromkeys(fehler):
            print(f"   - {f}")
        return 1
    print("Alle Pruefungen bestanden.")
    return 0


def main():
    global D

    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    if "--pruefen" in sys.argv:
        sys.exit(pruefen(content))

    try:
        D = _docx()
    except ImportError:
        print("!! python-docx ist nicht installiert — DOCX kann nicht gebaut werden.")
        print("   Die Textpruefung laeuft trotzdem:")
        print("       python Scripts/build_taschenbuch_docx_s2_1.py --pruefen")
        sys.exit(2)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"Lese: {INPUT_FILE}")

    doc = D['Document']()
    sec = doc.sections[0]
    setup_page(sec)
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    setup_mirror_margins(doc)
    create_styles(doc)
    add_front_matter(doc)

    doc.add_section(D['SEC'].NEW_PAGE)
    sec_body = doc.sections[-1]
    setup_page(sec_body)
    add_page_number_footer(sec_body)
    sec_body.header.is_linked_to_previous = False

    parse_and_build(doc, content)
    add_back_matter(doc)

    doc.save(OUTPUT_DOCX)
    print(f"DOCX erstellt: {OUTPUT_DOCX}")

    kapitel = len(re.findall(r'^# Kapitel ', content, re.MULTILINE))
    print(f"Kapitel: {kapitel}")
    if kapitel != ERWARTETE_KAPITEL:
        print(f"!! ABBRUCH: {kapitel} Kapitel, erwartet {ERWARTETE_KAPITEL}. Kein PDF.")
        sys.exit(1)

    print(f"Frontmatter aus build_manuskript_komplett_s2_1.py uebernommen:")
    print(f"   Widmung : {WIDMUNG_ZEILEN[0]}")
    print(f"   Epigraph: {EPIGRAPH_ZEILEN[0]}")

    if BAND_ASIN is None:
        print("⚠️  OHNE QR-CODE gebaut — keine ASIN hinterlegt.")
    print("⚠️  OHNE TEASER gebaut — S2-2 ist nicht geschrieben.")

    print("Konvertiere zu PDF...")
    convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)


if __name__ == "__main__":
    main()
