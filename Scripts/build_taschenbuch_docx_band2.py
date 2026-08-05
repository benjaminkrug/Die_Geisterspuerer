"""
Build a properly formatted DOCX + PDF for KDP Taschenbuch (6x9 inches)
from Manuskript_Band2_Komplett.md.
Handles: page size, mirror margins, page numbers, chapter headings,
paragraphs with first-line indent, scene breaks, italic/bold text,
front matter, back matter, Band 3 teaser.
Converts to PDF via Microsoft Word (win32com).
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Benjamin Krug"
BAND_NUM = 2
SERIES_TITLE = "Die Geisterspürer"
BAND_TITLE = "Der Friedhof ohne Namen"
BAND_SUBTITLE = f"{SERIES_TITLE} · Band {BAND_NUM}"
END_MARKER_PATTERN = r'\n---\n\n\*\*ENDE BAND 2\*\*\n\n---\n*'

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.join(_SCRIPT_DIR, "..")

INPUT_FILE = os.path.join(_ROOT, "Band2", "Manuskript", "Manuskript_Band2_Komplett.md")
OUTPUT_DIR = os.path.join(_ROOT, "Output", "Band2")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "KDP_Band2_Manuskript.docx")
OUTPUT_PDF = os.path.join(OUTPUT_DIR, "KDP_Band2_Manuskript.pdf")

SCENE_BREAK_SYMBOL = "\u2726  \u2726  \u2726"

# --- Amazon-Rezension: QR-Code auf der "Hat's dir gefallen?"-Seite ---
# ASIN des Taschenbuchs (vom Autor 2026-07-18). Bei neuer Ausgabe hier anpassen.
# Der Code wird von Scripts/build_qr_rezension.py erzeugt UND dort gegengelesen.
BAND_ASIN = "B0GV8R8QJ6"
REVIEW_URL = f"https://www.amazon.de/review/create-review?asin={BAND_ASIN}"
QR_IMAGE = os.path.join(_ROOT, "Band2", "Cover", "qr_rezension_band2.png")

Q = "\u201e"   # „
E = "\u201c"   # "


# ─────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────

def setup_page(section):
    """6x9 inch KDP paperback with mirror margins."""
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = Inches(0.875)   # inner (binding)
    section.right_margin = Inches(0.625)  # outer
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.35)


def setup_mirror_margins(doc):
    settings_element = doc.settings.element
    mirror = OxmlElement('w:mirrorMargins')
    settings_element.append(mirror)


# ─────────────────────────────────────────────
# FOOTER / PAGE NUMBERS
# ─────────────────────────────────────────────

def add_page_number_footer(section):
    """Centered page number in footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for fld_type, content in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        run = p.add_run()
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        if fld_type == 'instrText':
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = content
            run._r.append(instr)
        else:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), fld_type)
            run._r.append(fld)


# ─────────────────────────────────────────────
# STYLES
# ─────────────────────────────────────────────

def create_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_centered_image(doc, image_path, width_inches=1.6):
    """Fuegt ein zentriertes Bild ein (QR-Code). Fehlt die Datei, wird sie
    uebersprungen - dann steht im Buch aber KEIN Code. Deshalb meldet main()
    das am Ende ausdruecklich."""
    if not os.path.isfile(image_path):
        print(f"  WARNUNG: QR-Bild nicht gefunden, wird uebersprungen: {image_path}")
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    """Add a paragraph with inline *italic* and **bold** markdown formatting."""
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)

    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.findall(pattern, text)

    for full_match, bold_text, italic_text, plain_text in parts:
        if bold_text:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif italic_text:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif plain_text:
            run = p.add_run(plain_text)
            run.font.name = "Georgia"
            run.font.size = Pt(11)
    return p


def add_horizontal_rule(doc):
    """Thin centered line as section divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("── ◆ ──")
    run.font.name = "Georgia"
    run.font.size = Pt(10)
    run.font.color.rgb = None  # inherit


# ─────────────────────────────────────────────
# FRONT MATTER
# ─────────────────────────────────────────────

def add_front_matter(doc):
    # Half title
    add_blank_lines(doc, 8)
    add_centered_text(doc, SERIES_TITLE, font_size=20, bold=True)
    add_page_break(doc)

    # Full title page
    add_blank_lines(doc, 4)
    add_centered_text(doc, SERIES_TITLE, font_size=24, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, BAND_TITLE, font_size=17, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Band {BAND_NUM}", font_size=12)
    add_blank_lines(doc, 3)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # Copyright page
    add_blank_lines(doc, 10)
    add_centered_text(doc, f"{SERIES_TITLE} – {BAND_TITLE}", font_size=10, bold=True)
    add_centered_text(doc, f"Band {BAND_NUM}", font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = Pt(9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)


# ─────────────────────────────────────────────
# BACK MATTER
# ─────────────────────────────────────────────

def add_back_matter(doc):
    add_page_break(doc)

    # 1. Rezensions-Bitte
    add_blank_lines(doc, 3)
    add_centered_text(doc, f"{Q}{BAND_TITLE}{E} hat dir gefallen?", font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 2)

    # QR-Code direkt zum Bewertungsformular (identisch zu Band 1)
    add_centered_text(doc, "Einfach den Code scannen und eine Bewertung dalassen:",
                      font_size=11, italic=True)
    add_blank_lines(doc, 1)
    add_centered_image(doc, QR_IMAGE, width_inches=1.6)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "(Handykamera auf den Code halten \u2013 der Link \u00f6ffnet sich von selbst.)",
                      font_size=9, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # 2. Band 3 Teaser
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Weiterlesen? Hier kommt eine Vorschau auf Band 3:", font_size=11, italic=True)
    add_blank_lines(doc, 2)
    add_centered_text(doc, SERIES_TITLE, font_size=18, bold=True)
    add_centered_text(doc, "Schatten sieht mehr", font_size=14, italic=True)
    add_centered_text(doc, "Band 3", font_size=11)
    add_blank_lines(doc, 2)

    teaser_paragraphs = [
        "Man kann eine Nummer nicht anrufen, wenn niemand mehr da ist, der abhebt.",
        "Das wusste Nora noch nicht. Sie wählte trotzdem.",
        "Vor ihr auf dem Küchentisch lag die Visitenkarte. Drei Tage zuvor hatte Nora sie im Hausflur gefunden, die Tinte noch feucht — als hätte jemand sie eine Minute vorher geschrieben. Jetzt war sie trocken. Nora drehte sie zwischen den Fingern. Vorderseite: *M. Silber.* Eine Telefonnummer. Rückseite: zwei Worte.",
        "*Nicht alleine.*",
        "Sie hatte die Nummer schon sechsmal gewählt. Heute war das siebte Mal.",
        "Es klingelte.",
        "\"Sie geht nicht ran\", sagte Theo.",
        "Er saß ihr gegenüber, das Kinn auf den Armen, und sah dem Handy beim Klingeln zu, als wäre es ein Tier, das gleich beißen könnte.",
        "\"Vielleicht beim nächsten Mal.\"",
        "\"Das hast du beim vierten Mal auch gesagt.\"",
        "Es klingelte ein viertes Mal. Ein fünftes.",
        "Und dann — beim sechsten — ein Knacken. Ganz kurz. Als hätte am anderen Ende jemand den Hörer abgenommen. Nora drückte das Handy fester ans Ohr.",
        "\"Hallo?\", sagte sie. \"Frau Silber?\"",
        "Nichts. Kein Atmen, keine Stimme. Nur ein leises Rauschen, weit weg, wie Wind in einem leeren Raum. Dann das Tuten wieder, gleichmäßig und kalt.",
        "Beim siebten Klingeln kam ein langer Ton.",
        "Nora legte auf. Ihre Hand war nicht ganz ruhig.",
    ]

    for i, para in enumerate(teaser_paragraphs):
        add_formatted_paragraph(doc, para, first_line_indent=(i > 0))

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Jetzt überall erhältlich.", font_size=11, italic=True)

    # 3. Serienübersicht
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, f"{SERIES_TITLE} — Alle Bände", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    for num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(
            doc,
            f"**Band {num}:** {title}",
            first_line_indent=False,
        )

    # 4. Cross-Verweis
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Mehr spannende Abenteuer von {AUTHOR}:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schlüssel — und hinter der verschlossenen Tür wartet ein "
        "Geheimnis, das seit dreißig Jahren niemand lüften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Chrono-Agenten", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Drei Kinder. Ein Tablet, das die Zeit öffnet. Und ein Countdown, "
        "der nicht aufhört zu laufen.",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "Leo, Mila und Ben sind Agenten wider Willen — geschickt in die "
        "gefährlichsten Momente der Geschichte, um zu retten, was jemand "
        "zerstören will. Ihr Gegner heißt Codex. Und er kennt ihre Namen.",
        first_line_indent=False,
    )


# ─────────────────────────────────────────────
# CHAPTER HEADING + SCENE BREAK
# ─────────────────────────────────────────────

def add_chapter_heading(doc, title):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(36)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.name = "Georgia"
    run.bold = True


def add_scene_break(doc):
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)


# ─────────────────────────────────────────────
# PARSE + BUILD
# ─────────────────────────────────────────────

def parse_and_build(doc, content):
    # Find first chapter
    match = re.search(r'^# Kapitel 1', content, re.MULTILINE)
    if not match:
        raise ValueError("Konnte '# Kapitel 1' nicht finden.")
    body = content[match.start():]

    # Remove end marker
    body = re.sub(END_MARKER_PATTERN, '', body)
    body = re.sub(r'\n\*\*ENDE BAND 2\*\*\n*', '', body)

    lines = body.split('\n')

    def _naechste_zeile_ist_kapitel(start):
        """SEPARATOR-BUG-FIX (2026-07-18): schaut ueber Leerzeilen hinweg,
        ob als Naechstes eine Kapitelueberschrift kommt."""
        j = start
        while j < len(lines) and not lines[j].strip():
            j += 1
        return j < len(lines) and lines[j].strip().startswith('# Kapitel ')

    is_first_para_after_heading = False
    is_first_para_after_break = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith('# Kapitel '):
            title = line[2:]
            add_chapter_heading(doc, title)
            is_first_para_after_heading = True
            is_first_para_after_break = False
            i += 1
            continue

        if line == '---':
            # ── SEPARATOR-BUG-FIX (2026-07-18) ────────────────────────────
            # Die Kompilier-Skripte setzen ein "---" ZWISCHEN die Kapitel.
            # Fuer diesen Parser ist das nicht von einem Szenentrenner zu
            # unterscheiden -> im Druck stand ein Ornament hinter JEDEM
            # Cliffhanger, dann erst der Seitenumbruch (Band 1: 17, B2: 14,
            # B3: 15, B4: 15 Stueck).
            # Ein Szenentrenner direkt vor einer Kapitelueberschrift ist
            # IMMER falsch -> ueberspringen.
            if _naechste_zeile_ist_kapitel(i + 1):
                i += 1
                continue
            add_scene_break(doc)
            is_first_para_after_break = True
            is_first_para_after_heading = False
            i += 1
            continue

        no_indent = is_first_para_after_heading or is_first_para_after_break
        add_formatted_paragraph(doc, line, first_line_indent=not no_indent)
        is_first_para_after_heading = False
        is_first_para_after_break = False
        i += 1


# ─────────────────────────────────────────────
# PDF CONVERSION (via Word COM)
# ─────────────────────────────────────────────

def convert_to_pdf(docx_path, pdf_path):
    import subprocess
    soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
    out_dir = os.path.dirname(os.path.abspath(pdf_path))
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, os.path.abspath(docx_path)],
            check=True, capture_output=True
        )
        # LibreOffice names the output after the input file
        generated = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if generated != os.path.abspath(pdf_path) and os.path.exists(generated):
            os.replace(generated, os.path.abspath(pdf_path))
        print(f"PDF erstellt: {pdf_path}")
        return True
    except Exception as e:
        print(f"PDF-Konvertierung fehlgeschlagen: {e}")
        return False


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print(f"Lese: {INPUT_FILE}")
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Section 1: Front matter (no page numbers)
    section_front = doc.sections[0]
    setup_page(section_front)
    section_front.header.is_linked_to_previous = False
    section_front.footer.is_linked_to_previous = False
    setup_mirror_margins(doc)
    create_styles(doc)
    add_front_matter(doc)

    # Section 2: Body + back matter (with page numbers)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_body = doc.sections[-1]
    setup_page(section_body)
    add_page_number_footer(section_body)
    section_body.header.is_linked_to_previous = False

    parse_and_build(doc, content)
    add_back_matter(doc)

    doc.save(OUTPUT_DOCX)
    print(f"DOCX erstellt: {OUTPUT_DOCX}")

    chapters = len(re.findall(r'^# Kapitel ', content, re.MULTILINE))
    print(f"Kapitel: {chapters}")
    print(f"Zeichen: {len(content):,}")

    # Convert to PDF
    print("Konvertiere zu PDF...")
    convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)


if __name__ == "__main__":
    main()
