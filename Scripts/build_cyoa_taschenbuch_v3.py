"""
Erstellt ein KDP-fertiges Word-Dokument aus den CYOA v2 (Time Cave) Quelldateien.
Serie: "Die Geisterspürer"
Band 1: "Das Haus, das flüstert" (Entscheidungsbuch)

Visuelle Features (nach Herrenhaus-Vorbild):
- Drop Caps, gesperrte Überschriften, dekorative Linien
- Bookmarks + Hyperlinks + PAGEREF-Seitenzahlen
- Paragraph-Shading für Entscheidungen, Borders für Situationstexte
- Recto/Verso-Header, Mirror Margins
- Abschnitt-Verzeichnis mit Seitenzahlen
- Sackgassen + Enden: spezielles Styling
- Discovery-Seite, Notiz-Seiten, Serien-Uebersicht

Verwendung:
    python build_cyoa_taschenbuch_v3.py
"""
import re
import os
import random
import yaml
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# Konfiguration
# ============================================================

AUTHOR = "Benjamin Krug"

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.join(_SCRIPT_DIR, "..")
V2_DIR = os.path.join(_PROJECT_ROOT, "Band1", "CYOA", "v2")
GRAPH_FILE = os.path.join(V2_DIR, "graph_v2.yaml")
FRONTMATTER_FILE = os.path.join(_PROJECT_ROOT, "Band1", "CYOA", "frontmatter.md")
OUTPUT_DIR = os.path.join(_PROJECT_ROOT, "Output", "Band1")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "KDP_Band1_CYOA_v3_Manuskript.docx")

# Illustrationen (Platzhalter — Ordner wird erzeugt wenn nötig)
ILLUSTRATIONS_DIR = os.path.join(_PROJECT_ROOT, "Band1", "CYOA", "Illustrationen")
ILLUSTRATION_MAP = {
    "P1":   [2],      # Kirchgasse 14 — Das Haus
    "P2":   [3],      # Charakterbild: Nora, Theo, Schatten
    "P3":   [4],      # HILF am Fenster
    "A4":   [5],      # Der kalte Raum (Kratzer)
    "A11":  [6],      # Sie ist hier (Keller, allein)
    "A29":  [7],      # Die Karte (Artefakt Close-up)
    "A34":  [8],      # Graven am Friedhof
    "B4":   [9],      # Rote Tinte (Bibliothek)
    "B16":  [10],     # Der Schlüssel (Übergabe)
    "B21":  [11],     # Helds Warnung (blaues Licht)
    "C3":   [12],     # Theo rebelliert (Astronautenpyjama)
    "C5":   [13],     # Mama träumt (3 Uhr nachts)
    "C7e":  [14],     # Staffelübergabe (Tagebuch + Schlüssel)
    "E23":  [15],     # Mamas Frage (Pfannkuchen-Ende)
}
# Illustration 1 (Stadtkarte) wird als separate Seite in Front Matter eingefügt
ILLUSTRATION_WIDTH = Cm(8.0)

PAGE_WIDTH = Cm(12.7)    # 5 Zoll
PAGE_HEIGHT = Cm(20.32)  # 8 Zoll

# Deterministischer Seed für reproduzierbare Abschnitt-Nummerierung
RANDOM_SEED = 2026

# Globale Maps (werden in main() befüllt)
SECTION_MAP = {}   # Original-ID -> Neue Print-Nummer (str)
GRAPH_DATA = {}    # Geladener Graph

START_SECTION = "P1"


# ============================================================
# Graph laden
# ============================================================

def load_graph():
    """Lädt graph_v2.yaml und gibt das sections-Dict zurück."""
    with open(GRAPH_FILE, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return data.get("sections", {})


# ============================================================
# Abschnitt-Mapping (Scrambling)
# ============================================================

def create_section_mapping(sections):
    """Erstellt deterministisches Mapping: Original-ID -> Print-Nummer.

    P1 bleibt als 1 (wird in der Anleitung referenziert).
    Alle anderen bekommen neue, durchgemischte Nummern (2..N).
    """
    all_ids = sorted(sections.keys())
    other_ids = [sid for sid in all_ids if sid != START_SECTION]

    new_numbers = list(range(2, len(other_ids) + 2))
    rng = random.Random(RANDOM_SEED)
    rng.shuffle(new_numbers)

    mapping = {START_SECTION: "1"}
    for old_id, new_num in zip(other_ids, new_numbers):
        mapping[old_id] = str(new_num)

    return mapping


def map_id(old_id):
    """Gibt die gescramblte Print-Nummer für eine Section-ID zurück."""
    if not SECTION_MAP:
        return old_id
    result = SECTION_MAP.get(old_id)
    if result is None:
        # Führende Nullen entfernen: A02 -> A2, B01 -> B1
        stripped = re.sub(r'([A-Z]+)0+(\d)', r'\1\2', old_id)
        result = SECTION_MAP.get(stripped, old_id)
    return result


def bookmark_for(section_id):
    """Erzeugt einen Bookmark-Namen für eine Print-Nummer."""
    return f"abschnitt_{section_id}"


# ============================================================
# XML-Hilfsfunktionen (Bookmarks, Hyperlinks, PAGEREF)
# ============================================================

def add_bookmark(paragraph, name):
    """Fuegt einen Bookmark zu einem Absatz hinzu."""
    bm_id = str(abs(hash(name)) % 100000)
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bm_id)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bm_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _make_hyperlink_run(text, bold=False, italic=False, underline=True,
                        font_size=Pt(12)):
    """Erstellt ein w:r Element fuer einen Hyperlink."""
    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Georgia')
    rFonts.set(qn('w:hAnsi'), 'Georgia')
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)

    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    run_el.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run_el.append(t)
    return run_el


def add_hyperlink(paragraph, anchor, text, **kwargs):
    """Fuegt einen internen Hyperlink zu einem Absatz hinzu."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    hl.append(_make_hyperlink_run(text, **kwargs))
    paragraph._p.append(hl)


def add_paragraph_shading(paragraph, color="F0F0F0"):
    """Fuegt einem Absatz eine Hintergrund-Schattierung hinzu."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)


def add_paragraph_borders(paragraph, color="AAAAAA", size="4"):
    """Fuegt einem Absatz obere und untere Rahmenlinie hinzu."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), color)
        pBdr.append(border)
    pPr.append(pBdr)


def _add_pageref_field(paragraph, bookmark_name, font_size=Pt(9)):
    """Fuegt ein PAGEREF-Feld ein (zeigt die Seitenzahl eines Bookmarks)."""
    run1 = paragraph.add_run()
    run1.font.name = 'Georgia'
    run1.font.size = font_size
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fld_begin)

    run2 = paragraph.add_run()
    run2.font.name = 'Georgia'
    run2.font.size = font_size
    instr = OxmlElement('w:instrText')
    instr.text = f' PAGEREF {bookmark_name} '
    instr.set(qn('xml:space'), 'preserve')
    run2._r.append(instr)

    run3 = paragraph.add_run()
    run3.font.name = 'Georgia'
    run3.font.size = font_size
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld_end)


# ============================================================
# Illustrations-Support
# ============================================================

def get_illustration_path(number):
    """Gibt den Dateipfad fuer eine Illustrations-Nummer zurueck."""
    for pattern in [f"Illustration {number}.png", f"Illustration{number}.png",
                    f"Illustration {number}.jpg", f"Illustration{number}.jpg"]:
        path = os.path.join(ILLUSTRATIONS_DIR, pattern)
        if os.path.exists(path):
            return path
    return None


def add_illustration(doc, number):
    """Fuegt eine Illustration zentriert ins Dokument ein."""
    path = get_illustration_path(number)
    if not path:
        return
    doc.add_picture(path, width=ILLUSTRATION_WIDTH)
    pic_para = doc.paragraphs[-1]
    pic_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(6)
    pic_para.paragraph_format.space_after = Pt(6)


# ============================================================
# Dokument-Setup
# ============================================================

def setup_document():
    """Erstellt das Dokument mit KDP-Formatierung."""
    doc = Document()

    # Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.35
    pf.widow_control = True

    # Heading 1 Style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Georgia'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(14)
    h1.paragraph_format.page_break_before = False

    # Seitenformat
    sec = doc.sections[0]
    sec.page_width = PAGE_WIDTH
    sec.page_height = PAGE_HEIGHT
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(2.3)    # Bundsteg fuer Bindung
    sec.right_margin = Cm(1.3)

    # Mirror Margins + Recto/Verso
    doc.settings.element.append(OxmlElement('w:mirrorMargins'))
    doc.settings.element.append(OxmlElement('w:evenAndOddHeaders'))

    # Front Matter: Erste Seite ohne Header/Footer
    sec.different_first_page_header_footer = True

    # Front Matter: Leere Footer + Header
    footer = sec.footer
    footer.is_linked_to_previous = False
    header_odd = sec.header
    header_odd.is_linked_to_previous = False
    header_even = sec.even_page_header
    header_even.is_linked_to_previous = False

    return doc


# ============================================================
# Text-Rendering mit Markdown + Links + PAGEREF
# ============================================================

def _add_run(p, text, bold=False, italic=False):
    """Fuegt einen formatierten Run hinzu."""
    if not text:
        return
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True


def _render_text_with_links(p, text, bold=False, italic=False):
    """Rendert Text und macht 'Abschnitt X' / 'geh zu Abschnitt X' zu klickbaren Links.
    Ersetzt Original-IDs (A1, B16, etc.) durch Print-Nummern.
    Fuegt PAGEREF-Seitenzahlen hinzu."""

    # Erst alle Section-Referenzen finden und remappen
    def remap_match(m):
        prefix = m.group(1)
        old_id = m.group(2)
        new_id = map_id(old_id)
        return f"{prefix} {new_id}"

    text = re.sub(r'(Abschnitt|Ende)\s+([A-Z][A-Za-z_]*\d+[a-z]?)', remap_match, text)

    # Jetzt Links fuer "Abschnitt XX" einfuegen
    last = 0
    for m in re.finditer(r'Abschnitt\s+(\d+\w?)', text):
        before = text[last:m.start()]
        if before:
            _add_run(p, before, bold=bold, italic=italic)
        print_num = m.group(1)
        display_text = f"Abschnitt {print_num}"
        add_hyperlink(p, bookmark_for(print_num), display_text,
                      bold=bold, italic=italic)
        # Seitenzahl: " (S. XX)"
        _add_run(p, " (S.\u00a0", bold=bold, italic=italic)
        _add_pageref_field(p, bookmark_for(print_num), font_size=Pt(12))
        _add_run(p, ")", bold=bold, italic=italic)
        last = m.end()

    remaining = text[last:]
    if remaining:
        _add_run(p, remaining, bold=bold, italic=italic)


def render_markdown_line(p, text):
    """Rendert eine Zeile mit **bold**, *italic* und Abschnitt-Links."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _render_text_with_links(p, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*'):
            _render_text_with_links(p, part[1:-1], italic=True)
        else:
            _render_text_with_links(p, part)


# ============================================================
# Gruppen-Rendering (Drop Cap, Shading, Borders, Ornamente)
# ============================================================

def _render_group(doc, group, drop_cap=False, no_indent=False):
    """Rendert eine Gruppe von Zeilen als Absatz/Absaetze."""

    # Szenen-Trenner (---)
    if len(group) == 1 and group[0] == '---':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\u2726")
        run.font.size = Pt(12)
        run.font.name = 'Georgia'
        run.font.color.rgb = RGBColor(120, 120, 120)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        return

    # Navigation: "Weiter mit Abschnitt X"
    if len(group) == 1 and re.match(r'^\*Weiter\s+(mit|bei)\s+', group[0]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        nav_text = group[0].strip('*').strip()
        nav_text = "\u25B6 " + nav_text
        _render_text_with_links(p, nav_text, bold=True)
        for run in p.runs:
            run.font.size = Pt(13)
        add_paragraph_shading(p, "F0F0F0")
        return

    # ENDE-Marker
    if len(group) == 1 and group[0].strip() == '**ENDE**':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run("E N D E")
        run.font.name = 'Georgia'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        return

    # Einzelne Bold-Zeile (Entscheidung)
    if len(group) == 1 and group[0].startswith('**') and group[0].endswith('**'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        render_markdown_line(p, group[0])
        add_paragraph_shading(p, "F0F0F0")
        return

    # Einzelne Italic-Zeile (Situationstext / Entscheidungshinweis)
    if (len(group) == 1
            and group[0].startswith('*')
            and not group[0].startswith('**')
            and group[0].endswith('*')):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        render_markdown_line(p, group[0])
        add_paragraph_borders(p)
        return

    # Alle Zeilen Bold (Mehrfach-Entscheidungsblock)
    if all(l.startswith('**') and l.endswith('**') for l in group):
        for line in group:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            render_markdown_line(p, line)
            add_paragraph_shading(p, "F0F0F0")
        return

    # Alle Zeilen Italic (Ende-Hinweis / Epilog)
    if all(l.startswith('*') and not l.startswith('**') and l.endswith('*')
           for l in group):
        for line in group:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            render_markdown_line(p, line)
        return

    # Normaler Text: Zeilen zusammenfuegen
    joined = " ".join(group)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if no_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.paragraph_format.first_line_indent = Cm(0.5)

    if drop_cap and joined:
        first_char = joined[0]
        rest = joined[1:]
        dc_run = p.add_run(first_char)
        dc_run.font.name = 'Georgia'
        dc_run.font.size = Pt(28)
        dc_run.font.bold = True
        if rest:
            render_markdown_line(p, rest)
    else:
        render_markdown_line(p, joined)


# ============================================================
# Front Matter Seiten
# ============================================================

def add_half_title_page(doc):
    """Schmutztitelseite."""
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Das Haus, das fl\u00fcstert")
    run.font.size = Pt(16)
    run.font.name = 'Georgia'
    run.font.italic = True


def add_title_page(doc):
    """Titelseite."""
    doc.add_page_break()
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Serientitel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Die Geistersp\u00fcrer")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(12)

    # Dekorative Linie
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500\u2500\u2500\u2500\u2500\u2500\u2500")
    run.font.size = Pt(10)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(170, 170, 170)
    p.paragraph_format.space_after = Pt(12)

    # Buchtitel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Das Haus, das fl\u00fcstert")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(10)

    # Untertitel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ein interaktives Grusel-Abenteuer mit 24 Enden und einem geheimen Ende")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(6)

    # Statistik
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("3 Pfade. 24 Enden. 1 Geheimnis. 29 Entscheidungen.")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(10)

    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Autor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AUTHOR)
    run.font.size = Pt(14)
    run.font.name = 'Georgia'


def add_impressum_page(doc):
    """Impressum-Seite."""
    doc.add_page_break()
    for _ in range(2):
        doc.add_paragraph()

    lines = [
        "Die Geistersp\u00fcrer \u2014 Das Haus, das fl\u00fcstert",
        "Ein interaktives Grusel-Abenteuer mit 24 Enden und einem geheimen Ende",
        "",
        "\u00a9 2026 Benjamin Krug",
        "Alle Rechte vorbehalten.",
        "",
        "Independently published",
        "",
        "Dieses Buch ist ein Werk der Fiktion. Namen, Personen,",
        "Orte und Ereignisse sind frei erfunden. Jede \u00c4hnlichkeit",
        "mit tats\u00e4chlichen Personen, lebend oder verstorben,",
        "ist rein zuf\u00e4llig.",
    ]

    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if line:
            run = p.add_run(line)
            run.font.size = Pt(8)
            run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(2)


def add_dedication_page(doc):
    """Widmungsseite."""
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "F\u00fcr alle Kinder, die nachts lieber die Decke hochziehen\n"
        "\u2014 und trotzdem zuh\u00f6ren."
    )
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.name = 'Georgia'


def add_character_page(doc):
    """Charakter-Vorstellungsseite."""
    doc.add_page_break()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deine Begleiter")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(20)

    characters = [
        ("Nora (12)",
         "Zw\u00f6lf Jahre alt, logisch und aufmerksam. Nora stellt "
         "die richtigen Fragen \u2014 auch wenn sie die Antworten manchmal "
         "lieber nicht h\u00f6ren m\u00f6chte."),
        ("Theo (10)",
         "Noras kleiner Bruder. Hat vor allem Angst und macht trotzdem "
         "alles mit. Sein Humor ist seine R\u00fcstung \u2014 und manchmal "
         "seine st\u00e4rkste Waffe."),
        ("Schatten",
         "Ein schwarzer Hund, der eines Tages einfach da war. "
         "Er wei\u00df mehr, als ein Hund wissen sollte. Wenn sein Fell "
         "sich str\u00e4ubt, wird es unheimlich. Wenn er knurrt, wird es "
         "gef\u00e4hrlich."),
    ]

    for name, desc in characters:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(4)

        # Dekorative Mini-Linie
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\u2500\u2500\u2500")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(170, 170, 170)
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(16)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.right_indent = Cm(1.5)


def add_intro_page(doc):
    """Einleitungsseite aus frontmatter.md."""
    doc.add_page_break()

    if not os.path.exists(FRONTMATTER_FILE):
        print("  WARNUNG: frontmatter.md nicht gefunden!")
        return

    with open(FRONTMATTER_FILE, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    for line in lines:
        stripped = line.strip()

        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            continue

        if stripped.startswith('# '):
            title_text = stripped.lstrip('#').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'Georgia'
            p.paragraph_format.space_after = Pt(14)
            continue

        if stripped.startswith('## '):
            title_text = stripped.lstrip('#').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Georgia'
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(8)
            continue

        if stripped == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("\u2500\u2500\u2500")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(170, 170, 170)
            run.font.name = 'Georgia'
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        render_markdown_line(p, stripped)


def add_section_index(doc, total_sections):
    """Erstellt ein kompaktes Abschnitt-Verzeichnis mit Seitenzahlen."""
    doc.add_page_break()

    # Titel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abschnitt-Verzeichnis")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(4)

    # Hinweis
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Finde hier die Seite f\u00fcr jeden Abschnitt.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    # 5-Spalten-Tabelle ohne Rahmen
    cols = 5
    rows = (total_sections + cols - 1) // cols

    table = doc.add_table(rows=rows, cols=cols)

    # Rahmen entfernen
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)

    for i in range(total_sections):
        num = i + 1
        row_idx = i // cols
        col_idx = i % cols
        cell = table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        # Nummer (fett)
        run = p.add_run(f"{num:>2} ")
        run.font.name = 'Georgia'
        run.font.size = Pt(9)
        run.font.bold = True

        # Punkte
        run = p.add_run("\u00b7\u00b7\u00b7 ")
        run.font.name = 'Georgia'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(170, 170, 170)

        # Seitenzahl via PAGEREF
        _add_pageref_field(p, bookmark_for(str(num)), font_size=Pt(9))

    # Codewort-Hinweis
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run("Drei W\u00f6rter. Drei Pfade. Findest du sie?")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = 'Georgia'


def add_body_section_break(doc):
    """Fuegt Abschnittswechsel ein: Front Matter -> Body.

    Startet neue Word-Section mit:
    - Seitennummerierung ab 1
    - Footer mit Seitenzahl
    - Recto/Verso-Kopfzeilen
    """
    new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)

    # Seitenformat
    new_sec.page_width = PAGE_WIDTH
    new_sec.page_height = PAGE_HEIGHT
    new_sec.top_margin = Cm(1.6)
    new_sec.bottom_margin = Cm(1.6)
    new_sec.left_margin = Cm(2.3)
    new_sec.right_margin = Cm(1.3)

    # Seitennummerierung bei 1
    sectPr = new_sec._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    # Footer: zentrierte Seitenzahl
    footer = new_sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = fp.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fld_begin)

    run2 = fp.add_run()
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    instr.set(qn('xml:space'), 'preserve')
    run2._r.append(instr)

    run3 = fp.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld_end)

    for r in fp.runs:
        r.font.name = 'Georgia'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(170, 170, 170)

    # Ungerade Seiten (rechts): Serientitel
    header_odd = new_sec.header
    header_odd.is_linked_to_previous = False
    hp_odd = header_odd.paragraphs[0]
    hp_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp_odd.add_run("Die Geistersp\u00fcrer")
    run.font.name = 'Georgia'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(170, 170, 170)
    run.font.italic = True

    # Gerade Seiten (links): Buchtitel
    header_even = new_sec.even_page_header
    header_even.is_linked_to_previous = False
    hp_even = header_even.paragraphs[0]
    hp_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp_even.add_run("Das Haus, das fl\u00fcstert")
    run.font.name = 'Georgia'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(170, 170, 170)
    run.font.italic = True


# ============================================================
# Story-Abschnitte
# ============================================================

def add_story_section(doc, section_id, section_data, content):
    """Fuegt einen Story-Abschnitt zum Dokument hinzu.

    Styling variiert nach Typ:
    - story/choice: "A B S C H N I T T" + Nummer
    - dead_end: "S A C K G A S S E" + Nummer
    - ending: "E N D E" + Nummer + Titel
    """
    doc.add_page_break()

    sec_type = section_data.get("type", "story")
    title = section_data.get("title", "")
    print_num = map_id(section_id)

    # Leerraum oben
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Gesperrter Header-Text
    if sec_type == "dead_end":
        header_label = "S A C K G A S S E"
    elif sec_type == "ending":
        header_label = "E N D E"
    else:
        header_label = "A B S C H N I T T"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header_label)
    run.font.name = 'Georgia'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    p.paragraph_format.space_after = Pt(2)

    # Abschnitt-Nummer + Bookmark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(print_num)
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)
    add_bookmark(p, bookmark_for(print_num))

    # Dekorative Linie
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500\u2500\u2500")
    run.font.name = 'Georgia'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    p.paragraph_format.space_after = Pt(8)

    # Ending-Titel (kursiv, unter der Linie)
    if sec_type == "ending" and title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\u201e{title}\u201c")
        run.font.name = 'Georgia'
        run.font.size = Pt(13)
        run.font.italic = True
        p.paragraph_format.space_after = Pt(12)

    # Dead-End Titel
    if sec_type == "dead_end" and title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(120, 120, 120)
        p.paragraph_format.space_after = Pt(12)

    # Illustrationen einfuegen
    if section_id in ILLUSTRATION_MAP:
        for ill_num in ILLUSTRATION_MAP[section_id]:
            add_illustration(doc, ill_num)

    if not content:
        p = doc.add_paragraph()
        run = p.add_run("[Inhalt fehlt]")
        run.font.italic = True
        return

    # Content parsen
    lines = content.split('\n')

    # Gruppen bilden (getrennt durch Leerzeilen)
    groups = []
    current = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                groups.append(current)
                current = []
            continue
        if stripped.startswith('#'):
            continue
        current.append(stripped)
    if current:
        groups.append(current)

    # Gruppen rendern
    first_text_group = True
    after_separator = False

    for group in groups:
        is_sep = (len(group) == 1 and group[0] == '---')

        _render_group(doc, group,
                      drop_cap=(first_text_group and not is_sep),
                      no_indent=((first_text_group or after_separator)
                                 and not is_sep))

        if is_sep:
            after_separator = True
        else:
            if first_text_group:
                first_text_group = False
            after_separator = False

    # Dead-End: Rueckweg-Hinweis
    if sec_type == "dead_end":
        return_to = section_data.get("return_to", "")
        return_text = section_data.get("return_text", "")
        lesson = section_data.get("lesson", "")

        if return_to:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)

            # Ornament
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(8)
            p2.paragraph_format.space_after = Pt(8)
            run = p2.add_run("\u2500 \u25c6 \u2500")
            run.font.name = 'Georgia'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(170, 170, 170)

            # Lektion
            if lesson:
                p3 = doc.add_paragraph()
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p3.add_run(lesson)
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                run.font.italic = True
                p3.paragraph_format.space_after = Pt(8)

            # Zurueck-Link
            new_return_id = map_id(return_to)
            p4 = doc.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(4)
            p4.paragraph_format.space_after = Pt(4)

            _add_run(p4, "\u25B6 ", bold=True)
            text_label = return_text or f"Zur\u00fcck zu Abschnitt {new_return_id}"
            # Remap the return text
            text_label = re.sub(
                r'(Abschnitt)\s+([A-Z][A-Za-z_]*\d+[a-z]?)',
                lambda m: f"{m.group(1)} {map_id(m.group(2))}",
                text_label
            )
            add_hyperlink(p4, bookmark_for(new_return_id), text_label, bold=True)
            _add_run(p4, " (S.\u00a0", bold=True)
            _add_pageref_field(p4, bookmark_for(new_return_id), font_size=Pt(12))
            _add_run(p4, ")", bold=True)
            add_paragraph_shading(p4, "F0F0F0")
            return

    # Abschnitt-Ende-Ornament (nur fuer story/choice)
    if sec_type in ("story", "choice"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("\u2500 \u25c6 \u2500")
        run.font.name = 'Georgia'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(170, 170, 170)


# ============================================================
# Back Matter Seiten
# ============================================================

def add_discoveries_page(doc, sections):
    """Erstellt die 'Deine Entdeckungen' Sammelseite."""
    doc.add_page_break()

    # Titel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deine Entdeckungen")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    # -- Enden --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wie viele Enden hast du gefunden?")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(10)

    # Alle Endings aus dem Graph sortiert nach Print-Nummer
    endings = []
    for sid, data in sections.items():
        if data.get("type") == "ending":
            print_num = int(map_id(sid))
            title = data.get("title", sid)
            endings.append((print_num, sid, title))
    endings.sort()

    for print_num, sid, title in endings:
        p = doc.add_paragraph()
        run = p.add_run(f"\u2610  ENDE {print_num}")
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        run.font.small_caps = True
        run2 = p.add_run(f" \u2014 \u201e{title}\u201c")
        run2.font.size = Pt(10)
        run2.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.5)

    # Zaehler
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"_____ / {len(endings)} Enden gefunden!")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)

    # Trenner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2726")
    run.font.size = Pt(12)
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(10)

    # -- Codewort-Jagd --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Codewort-Jagd")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(8)

    codewords = [
        "Codewort 1: ________________",
        "Codewort 2: ________________",
        "Codewort 3: ________________",
    ]

    for cw in codewords:
        p = doc.add_paragraph()
        run = p.add_run(f"\u2610  {cw}")
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Tipp: Die W\u00f6rter sind GROSS geschrieben.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = 'Georgia'

    # Trenner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2726")
    run.font.size = Pt(12)
    run.font.name = 'Georgia'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)

    # -- Geheime Orte --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Geheime Orte")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(8)

    places = [
        "Silbers Wohnung entdeckt",
        "Linas Tagebuch gelesen",
        "Den Keller betreten",
        "Lina befreit",
        "Frau Held getroffen",
        "Den Friedhof besucht",
        "Graven gesehen",
        "Das geheime Ende gefunden",
    ]

    for place in places:
        p = doc.add_paragraph()
        run = p.add_run(f"\u2610  {place}")
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.5)


def add_notes_pages(doc):
    """Erstellt 2 Seiten Geisterspürer-Notizen."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deine Geistersp\u00fcrer-Notizen")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Schreib hier deine Hinweise und Entdeckungen auf!")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    for page in range(2):
        if page > 0:
            doc.add_page_break()
        for _ in range(20):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.right_indent = Cm(0.3)


def add_about_author_page(doc):
    """Über den Autor Seite."""
    doc.add_page_break()
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u00dcber den Autor")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    bio_lines = [
        ("Benjamin Krug schreibt Geschichten, die Kinder nicht mehr "
         "aus der Hand legen k\u00f6nnen \u2014 auch wenn sie sich "
         "vielleicht ein kleines bisschen f\u00fcrchten."),
        "",
        ("Wenn er nicht gerade an seinem Schreibtisch sitzt und sich "
         "neue Geheimnisse f\u00fcr Nora, Theo und Schatten ausdenkt, "
         "erkundet er am liebsten alte Geb\u00e4ude "
         "und fragt sich, welche Geschichten sie erz\u00e4hlen w\u00fcrden \u2014 "
         "wenn man nur genau genug zuh\u00f6rt."),
        "",
        ("Er lebt in Deutschland und glaubt fest daran, "
         "dass Zuh\u00f6ren die st\u00e4rkste Superkraft ist, "
         "die es gibt."),
    ]

    for line in bio_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line:
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)


def add_band2_teaser(doc):
    """Leseprobe Band 2."""
    doc.add_page_break()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Leseprobe")
    run.font.size = Pt(12)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(120, 120, 120)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Die Geistersp\u00fcrer")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Der Friedhof ohne Namen")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kapitel 1 \u2014 Zw\u00f6lf Punkte auf einer Karte")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    teaser_lines = [
        "Das Klopfen hatte aufgeh\u00f6rt.",
        "",
        "Seit drei Wochen. Seit Nora Lina befreit hatte. "
        "Seit die W\u00e4nde warm geworden waren und Schatten "
        "aufgeh\u00f6rt hatte, zur Treppe zu starren.",
        "",
        "Aber jetzt starrte er wieder.",
        "",
        "Nicht zur Treppe. Zum Fenster.",
        "",
        "\u00bbSchau mal\u00ab, sagte Theo. Er stand am K\u00fcchenfenster, "
        "die Nase fast an der Scheibe. \u00bbDer Friedhof.\u00ab",
        "",
        "Nora stellte sich neben ihn. "
        "Am Rand der Stadt, hinter den letzten H\u00e4usern, "
        "lag der alte Friedhof von Gravenstedt. "
        "Efeu \u00fcber den Mauern. Krumme Kreuze. "
        "Und ein blaues Licht, das zwischen den Grabsteinen pulsierte.",
        "",
        "\u00bbDas ist nicht normal\u00ab, sagte Theo.",
        "",
        "\u00bbNein\u00ab, sagte Nora. \u00bbDas ist Nummer zwei.\u00ab",
    ]

    for line in teaser_lines:
        p = doc.add_paragraph()
        if line:
            run = p.add_run(line)
            run.font.name = 'Georgia'
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2726")
    run.font.size = Pt(12)
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Band 2: Der Friedhof ohne Namen")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bald erh\u00e4ltlich!")
    run.font.size = Pt(12)
    run.font.name = 'Georgia'
    run.font.bold = True


def add_series_overview_page(doc):
    """Serienübersicht."""
    doc.add_page_break()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mehr Abenteuer von Benjamin Krug")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(20)

    series = [
        ("Die Geistersp\u00fcrer",
         "Zwei Geschwister. Ein Hund, der mehr wei\u00df als er sollte. "
         "Und eine Stadt voller Geister, die darauf warten, geh\u00f6rt zu werden.",
         ["Band 1: Das Haus, das fl\u00fcstert",
          "Band 1: Das Haus, das fl\u00fcstert \u2014 Entscheidungsbuch",
          "Band 2: Der Friedhof ohne Namen",
          "Band 3: Schatten sieht mehr",
          "Band 4: Die zugemauerte T\u00fcr",
          "Band 5: Der Schleier"]),
        ("Die Herrenhaus-Detektive",
         "Niemand darf das alte Herrenhaus betreten. "
         "Aber Jonas, Mila und Ben finden einen Schl\u00fcssel \u2014 "
         "und hinter der verschlossenen T\u00fcr wartet ein Geheimnis.",
         ["Band 1: Das verbotene Herrenhaus",
          "Band 1: Das verbotene Herrenhaus \u2014 Entscheidungsbuch",
          "Band 2: Das Geheimnis des Brunnens"]),
    ]

    for i, (title, desc, bands) in enumerate(series):
        if i > 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("\u2014 \u2014 \u2014")
            run.font.size = Pt(10)
            run.font.name = 'Georgia'
            run.font.color.rgb = RGBColor(120, 120, 120)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(6)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        run.font.italic = True
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)

        for band in bands:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(band)
            run.font.size = Pt(10)
            run.font.name = 'Georgia'
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Alle B\u00fccher auf Amazon erh\u00e4ltlich!")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Georgia'


def add_review_request_page(doc):
    """Rezensions-Bitte."""
    doc.add_page_break()
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Eine kleine Bitte")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2605 \u2605 \u2605 \u2605 \u2605")
    run.font.size = Pt(14)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(200, 160, 50)
    p.paragraph_format.space_after = Pt(20)

    review_lines = [
        ("Wenn dir das Buch gefallen hat, w\u00fcrde ich mich riesig "
         "\u00fcber eine kurze Bewertung auf Amazon freuen."),
        "",
        ("Ein paar Worte reichen v\u00f6llig \u2014 zum Beispiel, welches "
         "Ende dir am besten gefallen hat oder ob Theo dich zum Lachen "
         "gebracht hat."),
        "",
        ("Jede Bewertung hilft mir, noch mehr Geistersp\u00fcrer-Abenteuer "
         "f\u00fcr dich zu schreiben!"),
    ]

    for line in review_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line:
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Danke! \u2764")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2014 Benjamin")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.name = 'Georgia'


def add_acknowledgments_page(doc):
    """Danksagung."""
    doc.add_page_break()
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Danksagung")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    p.paragraph_format.space_after = Pt(14)

    thanks_text = (
        "Danke an alle kleinen und gro\u00dfen Testleser, "
        "die dieses Buch besser gemacht haben. "
        "Und an dich \u2014 weil du zugeh\u00f6rt hast, "
        "als Lina es am meisten brauchte."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(thanks_text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.name = 'Georgia'
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)


# ============================================================
# Content laden
# ============================================================

def load_section_content(file_path):
    """Lädt eine Markdown-Section-Datei und gibt den Body-Text zurück."""
    full_path = os.path.join(V2_DIR, file_path)
    if not os.path.exists(full_path):
        return None

    with open(full_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Heading entfernen
    content = re.sub(r'^#[^\n]*\n+', '', content)
    content = re.sub(r'^##[^\n]*\n+', '', content, flags=re.MULTILINE)

    return content.strip()


# ============================================================
# Main
# ============================================================

def main():
    global SECTION_MAP, GRAPH_DATA

    print("Erstelle Geisterspürer CYOA v3 Manuskript...")

    # Graph laden
    print("  Lade graph_v2.yaml...")
    sections = load_graph()
    GRAPH_DATA = sections
    all_ids = sorted(sections.keys())
    print(f"  Gefunden: {len(all_ids)} Abschnitte")

    # Mapping erstellen
    print("  Erstelle Abschnitt-Mapping...")
    SECTION_MAP = create_section_mapping(sections)

    # Reihenfolge: P1 zuerst, Rest nach Print-Nummer sortiert
    ordered_ids = sorted(all_ids, key=lambda sid: int(map_id(sid)))

    print(f"  P1 -> Abschnitt 1")
    print(f"  {len(ordered_ids)} Abschnitte gesamt")

    # Content laden
    print("  Lade Inhalte...")
    contents = {}
    missing = []
    for sid in all_ids:
        sec = sections[sid]
        file_path = sec.get("file")
        if file_path:
            content = load_section_content(file_path)
            if content is None:
                missing.append(sid)
            contents[sid] = content
        else:
            missing.append(sid)
            contents[sid] = None

    if missing:
        print(f"  WARNUNG: {len(missing)} Abschnitte ohne Datei: {missing[:10]}")

    # Output-Verzeichnis
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Dokument erstellen
    print("  Erstelle DOCX...")
    doc = setup_document()

    # === FRONT MATTER ===
    print("  Schmutztitel")
    add_half_title_page(doc)

    print("  Titelseite")
    add_title_page(doc)

    print("  Impressum")
    add_impressum_page(doc)

    print("  Widmung")
    add_dedication_page(doc)

    print("  Charaktere")
    add_character_page(doc)

    # Stadtkarte Gravenstedt (Illustration 1, separate Seite)
    print("  Stadtkarte Gravenstedt")
    add_illustration(doc, 1)

    print("  Anleitung (frontmatter.md)")
    add_intro_page(doc)

    print("  Abschnitt-Verzeichnis")
    add_section_index(doc, len(ordered_ids))

    # === SECTION BREAK ===
    print("  Section Break (Seitennummerierung ab 1)")
    add_body_section_break(doc)

    # === BODY ===
    total_words = 0
    for sid in ordered_ids:
        sec = sections[sid]
        content = contents.get(sid)
        print_num = map_id(sid)
        sec_type = sec.get("type", "story")
        print(f"  [{sec_type:>8s}] {sid:>6s} -> Abschnitt {print_num}")
        add_story_section(doc, sid, sec, content)

        if content:
            total_words += len(content.split())

    # === BACK MATTER ===
    print("  Deine Entdeckungen")
    add_discoveries_page(doc, sections)

    print("  Geisterspürer-Notizen")
    add_notes_pages(doc)

    print("  Über den Autor")
    add_about_author_page(doc)

    print("  Band 2 Leseprobe")
    add_band2_teaser(doc)

    print("  Serienübersicht")
    add_series_overview_page(doc)

    print("  Rezensions-Bitte")
    add_review_request_page(doc)

    print("  Danksagung")
    add_acknowledgments_page(doc)

    # Speichern
    doc.save(OUTPUT_FILE)
    print(f"\nManuskript erstellt: {OUTPUT_FILE}")
    print(f"  {len(ordered_ids)} Abschnitte")
    print(f"  ~{total_words:,} Wörter gesamt")
    print(f"  Fehlende Dateien: {len(missing)}")

    # Mapping ausgeben (erste 15)
    print("\n--- Abschnitt-Zuordnung (intern -> Druck) ---")
    for sid in ordered_ids[:15]:
        sec = sections[sid]
        sec_type = sec.get("type", "story")
        print(f"  {sid:>6s} -> {map_id(sid):>3s}  [{sec_type}]")
    if len(ordered_ids) > 15:
        print(f"  ... ({len(ordered_ids) - 15} weitere)")


if __name__ == "__main__":
    main()
