"""
Baut ein KDP-fertiges DOCX + PDF (Taschenbuch 6x9 Zoll) fuer Band 3
aus Manuskript_Band3_Komplett.md.

Erledigt: Seitengroesse, Spiegelraender, Seitenzahlen, Kapitelueberschriften,
Absaetze mit Erstzeileneinzug, Szenentrenner, kursiv/fett,
Front Matter (inkl. WIDMUNG + EPIGRAPH), Back Matter, Band-4-Teaser.
PDF-Konvertierung via LibreOffice (soffice) oder Word.

Verwendung:
    python build_taschenbuch_docx_band3.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Benjamin Krug"
BAND_NUM = 3
SERIES_TITLE = "Die Geisterspürer"
BAND_TITLE = "Schatten sieht mehr"
BAND_SUBTITLE = f"{SERIES_TITLE} · Band {BAND_NUM}"

# Naechster Band (fuer den Teaser am Ende)
NEXT_BAND_NUM = 4
NEXT_BAND_TITLE = "Die zugemauerte Tür"

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.join(_SCRIPT_DIR, "..")

INPUT_FILE = os.path.join(_ROOT, "Band3", "Manuskript", "Manuskript_Band3_Komplett.md")
OUTPUT_DIR = os.path.join(_ROOT, "Output", "Band3")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "KDP_Band3_Manuskript.docx")
OUTPUT_PDF = os.path.join(OUTPUT_DIR, "KDP_Band3_Manuskript.pdf")

SCENE_BREAK_SYMBOL = "✦  ✦  ✦"

# --- Amazon-Rezension: QR-Code auf der "Hat's dir gefallen?"-Seite ---
# ASIN des Taschenbuchs (vom Autor 2026-07-18). Bei neuer Ausgabe hier anpassen.
# Der Code wird von Scripts/build_qr_rezension.py erzeugt UND dort gegengelesen.
BAND_ASIN = "B0H4VQHBLX"
REVIEW_URL = f"https://www.amazon.de/review/create-review?asin={BAND_ASIN}"
QR_IMAGE = os.path.join(_ROOT, "Band3", "Cover", "qr_rezension_band3.png")

Q = "„"   # „
E = "“"   # "

# Anzahl Woerter am Kapitelanfang, die in Kapitaelchen gesetzt werden
SMALLCAPS_WORDS = 4


# ─────────────────────────────────────────────
# TYPOGRAFIE  (deutsche Anfuehrungszeichen, Apostroph, Gedankenstrich)
# ─────────────────────────────────────────────

def typo_quotes(text: str) -> str:
    """Wandelt gerade ASCII-Anfuehrungszeichen in deutsche „…" um.

    Heuristik (paarweise, robust): Wir tracken, ob wir uns INNERHALB einer
    Anfuehrung befinden. Das erste " oeffnet, das naechste schliesst, usw.
    Das ist fuer Romandialoge zuverlaessiger als reine Nachbarschafts-Regeln
    (loest korrekt: abgebrochene Rede „Schatten—", kursive Markdown-Echos
    *"Mama?"*, mehrere Dialoge pro Absatz). Apostroph ' -> '."""
    out = []
    inside = False  # sind wir gerade zwischen einem oeffnenden und schliessenden "?
    for idx, ch in enumerate(text):
        if ch == '"':
            if not inside:
                out.append('„')      # oeffnend
                inside = True
            else:
                out.append('“')      # schliessend
                inside = False
        elif ch == "'":
            # Apostroph nur als typografischen ' setzen, wenn er INNERHALB
            # eines Wortes steht (z.B. "geht's", "hab's"). Sonst unveraendert.
            prev = text[idx - 1] if idx > 0 else ' '
            nxt = text[idx + 1] if idx + 1 < len(text) else ' '
            if prev.isalpha() and nxt.isalpha():
                out.append('’')      # '
            elif prev.isalpha():
                out.append('’')      # geht's. -> geht's
            else:
                out.append(ch)
        else:
            out.append(ch)
    return "".join(out)


def typo_dashes(text: str) -> str:
    """Vereinheitlicht Gedankenstriche auf den deutschen Halbgeviertstrich (–)
    mit umgebenden Leerzeichen (wie Band 1). Wandelt Geviertstrich (—) um,
    NUR wenn er als Gedankenstrich (von Leerzeichen umgeben oder am Wortende)
    steht. Bindestriche in Woertern bleiben unberuehrt."""
    # "— " / " —" / " — " (Geviert als Gedankenstrich) -> Halbgeviert
    text = re.sub(r'\s—\s', ' – ', text)      #  —  ->  –
    text = re.sub(r'\s—', ' –', text)          #  —  am Wortanfang
    text = re.sub(r'—\s', '– ', text)          # —  am Wortende (z.B. "nie—")
    # alleinstehender Geviert ohne Leerzeichen (selten): vorsichtig lassen
    return text


def apply_typography(text: str) -> str:
    """Alle typografischen Verbesserungen auf einen Textabsatz anwenden."""
    text = typo_quotes(text)
    text = typo_dashes(text)
    return text

# ── Literarische Frontmatter (gewaehlt mit dem Autor, 2026-06-10) ──────────────
WIDMUNG_ZEILEN = [
    "Für alle, denen einmal jemand richtig zugehört hat.",
    "Und für die, die noch darauf warten.",
]
EPIGRAPH_ZEILEN = [
    "Nicht jeder Unruhige sucht die Wahrheit.",
    "Manche suchen nur die Erlaubnis, endlich aufzuhören.",
]
EPIGRAPH_SOURCE = "aus dem Notizbuch von Margret Silber"


# ─────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────

def setup_page(section):
    """6x9 inch KDP paperback with mirror margins."""
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = Inches(0.875)   # inner (binding)
    section.right_margin = Inches(0.625)  # outer
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.35)


def setup_mirror_margins(doc):
    settings_element = doc.settings.element
    mirror = OxmlElement('w:mirrorMargins')
    settings_element.append(mirror)


# ─────────────────────────────────────────────
# FOOTER / PAGE NUMBERS
# ─────────────────────────────────────────────

def add_page_number_footer(section):
    """Centered page number in footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for fld_type, content in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        run = p.add_run()
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        if fld_type == 'instrText':
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = content
            run._r.append(instr)
        else:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), fld_type)
            run._r.append(fld)


# ─────────────────────────────────────────────
# STYLES
# ─────────────────────────────────────────────

def create_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_centered_image(doc, image_path, width_inches=1.6):
    """Fuegt ein zentriertes Bild ein (QR-Code). Fehlt die Datei, wird sie
    uebersprungen - dann steht im Buch aber KEIN Code. Deshalb meldet main()
    das am Ende ausdruecklich."""
    if not os.path.isfile(image_path):
        print(f"  WARNUNG: QR-Bild nicht gefunden, wird uebersprungen: {image_path}")
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    """Add a paragraph with inline *italic* and **bold** markdown formatting.
    Typografie (deutsche Quotes, Apostroph, Striche) wird automatisch
    angewandt — idempotent, also auch fuer bereits gesetzte Backmatter sicher."""
    text = apply_typography(text)
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)

    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.findall(pattern, text)

    for full_match, bold_text, italic_text, plain_text in parts:
        if bold_text:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif italic_text:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif plain_text:
            run = p.add_run(plain_text)
            run.font.name = "Georgia"
            run.font.size = Pt(11)
    return p


# ─────────────────────────────────────────────
# FRONT MATTER  (Halbtitel · Titel · Impressum · WIDMUNG · EPIGRAPH)
# ─────────────────────────────────────────────

def add_front_matter(doc):
    # Half title
    add_blank_lines(doc, 8)
    add_centered_text(doc, SERIES_TITLE, font_size=20, bold=True)
    add_page_break(doc)

    # Full title page
    add_blank_lines(doc, 4)
    add_centered_text(doc, SERIES_TITLE, font_size=24, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, BAND_TITLE, font_size=17, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Band {BAND_NUM}", font_size=12)
    add_blank_lines(doc, 3)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # Copyright page
    add_blank_lines(doc, 10)
    add_centered_text(doc, f"{SERIES_TITLE} – {BAND_TITLE}", font_size=10, bold=True)
    add_centered_text(doc, f"Band {BAND_NUM}", font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = Pt(9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)

    # ── WIDMUNG (eigene Seite, zentriert, kursiv) ──────────────────────────
    add_page_break(doc)
    add_blank_lines(doc, 12)
    for zeile in WIDMUNG_ZEILEN:
        add_centered_text(doc, zeile, font_size=12, italic=True)
        add_blank_lines(doc, 1)

    # ── EPIGRAPH (eigene Seite, Silber-Notizbuch) ──────────────────────────
    add_page_break(doc)
    add_blank_lines(doc, 12)
    for zeile in EPIGRAPH_ZEILEN:
        add_centered_text(doc, zeile, font_size=12, italic=True)
        add_blank_lines(doc, 1)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"— {EPIGRAPH_SOURCE}", font_size=10, italic=True)


# ─────────────────────────────────────────────
# BACK MATTER
# ─────────────────────────────────────────────

def add_back_matter(doc):
    add_page_break(doc)

    # 1. Rezensions-Bitte
    add_blank_lines(doc, 3)
    add_centered_text(doc, f"{Q}{BAND_TITLE}{E} hat dir gefallen?", font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 2)

    # QR-Code direkt zum Bewertungsformular (identisch zu Band 1)
    add_centered_text(doc, "Einfach den Code scannen und eine Bewertung dalassen:",
                      font_size=11, italic=True)
    add_blank_lines(doc, 1)
    add_centered_image(doc, QR_IMAGE, width_inches=1.6)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "(Handykamera auf den Code halten \u2013 der Link \u00f6ffnet sich von selbst.)",
                      font_size=9, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # 2. Band 4 Teaser (Auszug passend zur Serien-Fortsetzung)
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(
        doc, f"Weiterlesen? Hier kommt eine Vorschau auf Band {NEXT_BAND_NUM}:",
        font_size=11, italic=True,
    )
    add_blank_lines(doc, 2)
    add_centered_text(doc, SERIES_TITLE, font_size=18, bold=True)
    add_centered_text(doc, NEXT_BAND_TITLE, font_size=14, italic=True)
    add_centered_text(doc, f"Band {NEXT_BAND_NUM}", font_size=11)
    add_blank_lines(doc, 2)

    teaser_paragraphs = [
        "Manche Türen sind schon offen, lange bevor man sie findet.",
        "Nora wusste das noch nicht. Sie saß am Küchentisch und suchte auf einer Karte nach etwas, das über ihr lag.",
        "Vor ihr die alte Karte von Gravenstedt. Frau Silbers Karte, handgezeichnet, mit den zwölf Markierungen. Drei davon durchgestrichen jetzt.",
        "Lina. Brenner. Marlene.",
        "Neun blieben.",
        "Aber das war nicht das, was Nora anstarrte.",
        "Es war der rote Kreis am Rand. Der, in dem seit dem ersten Tag nur ein Wort stand.",
        "*Graven.*",
        "Über Nacht war darin ein Zeichen erschienen. Klein, dunkel, scharf. Niemand hatte es geschrieben.",
        "\"Es ist immer noch da\", sagte Theo.",
        "Er saß ihr gegenüber und drehte etwas zwischen den Fingern. Einen flachen, glatten Stein, nicht größer als eine Münze, mit einem Loch in der Mitte. Er tat das immer, wenn ihm etwas nicht geheuer war.",
        "\"Ich hab gehofft, es wär weg\", sagte er. \"Wenn man morgens aufwacht und der Albtraum ist verschwunden.\"",
        "\"Albträume verschwinden nicht, weil man es hofft.\" Nora sagte es ruhig, fast beiläufig. \"Und das hier ist kein Albtraum. Es ist Tinte.\"",
        "\"Tinte, die sich selbst schreibt.\" Theo schob den Stein in die Jackentasche. \"Das ist die gruseligste Art von Tinte.\"",
        "Nora hielt die Karte gegen das Fenster. Draußen war grauer Morgen, das Licht fiel schräg durchs Papier. Sie hatte gedacht, das Zeichen wäre einfach ein Kritzel. Ein Fleck.",
        "Es war kein Fleck.",
        "Im Licht sah sie es genauer. Feine Striche, alle gerade, alle im Winkel. Kein Symbol.",
        "Eine Zeichnung.",
    ]

    for i, para in enumerate(teaser_paragraphs):
        add_formatted_paragraph(doc, para, first_line_indent=(i > 0))

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Jetzt überall erhältlich.", font_size=11, italic=True)

    # 3. Serienübersicht
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, f"{SERIES_TITLE} — Alle Bände", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    for num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(
            doc,
            f"**Band {num}:** {title}",
            first_line_indent=False,
        )

    # 4. Cross-Verweis
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Mehr spannende Abenteuer von {AUTHOR}:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schlüssel — und hinter der verschlossenen Tür wartet ein "
        "Geheimnis, das seit dreißig Jahren niemand lüften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Chrono-Agenten", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Drei Kinder. Ein Tablet, das die Zeit öffnet. Und ein Countdown, "
        "der nicht aufhört zu laufen.",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "Leo, Mila und Ben sind Agenten wider Willen — geschickt in die "
        "gefährlichsten Momente der Geschichte, um zu retten, was jemand "
        "zerstören will. Ihr Gegner heißt Codex. Und er kennt ihre Namen.",
        first_line_indent=False,
    )


# ─────────────────────────────────────────────
# CHAPTER HEADING + SCENE BREAK
# ─────────────────────────────────────────────

def add_chapter_heading(doc, title):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(36)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.name = "Georgia"
    run.bold = True


def add_scene_break(doc):
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)


def add_chapter_first_paragraph(doc, text):
    """Erster Absatz nach einer Kapitelueberschrift: die ersten SMALLCAPS_WORDS
    Woerter werden in Kapitaelchen (small caps) gesetzt — eleganter Kapitelauftakt
    wie bei vielen Mittelgrad-Bestsellern. Kein Einzug (Kapitelanfang).
    Faellt auf normalen Absatz zurueck, wenn der Anfang Sonderformatierung
    (z.B. kursives Cliffhanger-Echo *...* ) traegt."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

    # Wenn der Absatz mit kursivem/fettem Markdown beginnt, KEINE Kapitaelchen
    # (z.B. Kap. 3 "*Wo habt ihr den her?*", Kap. 13 "*Du bist nicht Jonas.*")
    # -> dann normaler formatierter Absatz, damit die Kursivierung erhalten bleibt.
    if text.lstrip().startswith('*'):
        _add_runs_with_markdown(p, text)
        return p

    # Trenne: erste N "Woerter" als Kapitaelchen, Rest normal.
    # Ein "Wort" = Sequenz Nicht-Leerzeichen; wir zaehlen bis SMALLCAPS_WORDS.
    m = re.match(r'^(\s*)(.*)$', text)
    leading_ws, rest = m.group(1), m.group(2)

    tokens = rest.split(' ')
    head_tokens = tokens[:SMALLCAPS_WORDS]
    tail_tokens = tokens[SMALLCAPS_WORDS:]
    head = ' '.join(head_tokens)
    tail = (' ' + ' '.join(tail_tokens)) if tail_tokens else ''

    # Falls der Kopf Markdown-Sterne enthaelt, lieber sicher normal rendern.
    if '*' in head:
        _add_runs_with_markdown(p, text)
        return p

    if leading_ws:
        run = p.add_run(leading_ws)
        run.font.name = "Georgia"
        run.font.size = Pt(11)

    # Kapitaelchen-Run
    run = p.add_run(head)
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.font.small_caps = True

    # Rest mit normaler Markdown-Formatierung
    if tail:
        _add_runs_with_markdown(p, tail)
    return p


def _add_runs_with_markdown(p, text):
    """Hilfsfunktion: haengt Runs mit *italic*/**bold**-Markdown an Paragraph p."""
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for full_match, bold_text, italic_text, plain_text in re.findall(pattern, text):
        if bold_text:
            run = p.add_run(bold_text); run.bold = True
        elif italic_text:
            run = p.add_run(italic_text); run.italic = True
        elif plain_text:
            run = p.add_run(plain_text)
        else:
            continue
        run.font.name = "Georgia"
        run.font.size = Pt(11)


# ─────────────────────────────────────────────
# PARSE + BUILD
# ─────────────────────────────────────────────

def parse_and_build(doc, content):
    # Find first chapter (alles davor = Frontmatter aus dem MD, wird ignoriert)
    match = re.search(r'^# Kapitel 1', content, re.MULTILINE)
    if not match:
        raise ValueError("Konnte '# Kapitel 1' nicht finden.")
    body = content[match.start():]

    # Remove end marker
    body = re.sub(r'\n---\n\n\*\*ENDE BAND 3\*\*\n\n---\n*', '', body)
    body = re.sub(r'\n\*\*ENDE BAND 3\*\*\n*', '', body)

    lines = body.split('\n')

    def _naechste_zeile_ist_kapitel(start):
        """SEPARATOR-BUG-FIX (2026-07-18): schaut ueber Leerzeilen hinweg,
        ob als Naechstes eine Kapitelueberschrift kommt."""
        j = start
        while j < len(lines) and not lines[j].strip():
            j += 1
        return j < len(lines) and lines[j].strip().startswith('# Kapitel ')

    is_first_para_after_heading = False
    is_first_para_after_break = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith('# Kapitel '):
            title = line[2:]
            add_chapter_heading(doc, title)
            is_first_para_after_heading = True
            is_first_para_after_break = False
            i += 1
            continue

        if line == '---':
            # ── SEPARATOR-BUG-FIX (2026-07-18) ────────────────────────────
            # Die Kompilier-Skripte setzen ein "---" ZWISCHEN die Kapitel.
            # Fuer diesen Parser ist das nicht von einem Szenentrenner zu
            # unterscheiden -> im Druck stand ein Ornament hinter JEDEM
            # Cliffhanger, dann erst der Seitenumbruch (Band 1: 17, B2: 14,
            # B3: 15, B4: 15 Stueck).
            # Ein Szenentrenner direkt vor einer Kapitelueberschrift ist
            # IMMER falsch -> ueberspringen.
            if _naechste_zeile_ist_kapitel(i + 1):
                i += 1
                continue
            add_scene_break(doc)
            is_first_para_after_break = True
            is_first_para_after_heading = False
            i += 1
            continue

        no_indent = is_first_para_after_heading or is_first_para_after_break
        # Typografie: gerade Quotes -> „…", Apostroph, einheitliche Striche
        text = apply_typography(line)
        # Kapitaelchen fuer die ersten Woerter direkt nach einer Kapitelueberschrift
        if is_first_para_after_heading:
            add_chapter_first_paragraph(doc, text)
        else:
            add_formatted_paragraph(doc, text, first_line_indent=not no_indent)
        is_first_para_after_heading = False
        is_first_para_after_break = False
        i += 1


# ─────────────────────────────────────────────
# PDF CONVERSION (LibreOffice, sonst Word COM als Fallback)
# ─────────────────────────────────────────────

def convert_to_pdf(docx_path, pdf_path):
    import subprocess
    out_dir = os.path.dirname(os.path.abspath(pdf_path))

    # 1) Versuch: LibreOffice (soffice)
    for soffice in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
    ):
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, os.path.abspath(docx_path)],
                check=True, capture_output=True,
            )
            generated = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if os.path.exists(generated):
                if os.path.abspath(generated) != os.path.abspath(pdf_path):
                    os.replace(generated, os.path.abspath(pdf_path))
                print(f"PDF erstellt (LibreOffice): {pdf_path}")
                return True
        except Exception:
            continue

    # 2) Fallback: Microsoft Word via COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF
        doc.Close()
        word.Quit()
        print(f"PDF erstellt (Word): {pdf_path}")
        return True
    except Exception as e:
        print(f"PDF-Konvertierung fehlgeschlagen (LibreOffice + Word nicht verfuegbar): {e}")
        print(f"  -> DOCX ist fertig: {docx_path}")
        print(f"  -> PDF manuell aus Word/LibreOffice exportieren (Datei > Exportieren als PDF).")
        return False


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print(f"Lese: {INPUT_FILE}")
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Section 1: Front matter (no page numbers)
    section_front = doc.sections[0]
    setup_page(section_front)
    section_front.header.is_linked_to_previous = False
    section_front.footer.is_linked_to_previous = False
    setup_mirror_margins(doc)
    create_styles(doc)
    add_front_matter(doc)

    # Section 2: Body + back matter (with page numbers)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_body = doc.sections[-1]
    setup_page(section_body)
    add_page_number_footer(section_body)
    section_body.header.is_linked_to_previous = False

    parse_and_build(doc, content)
    add_back_matter(doc)

    doc.save(OUTPUT_DOCX)
    print(f"DOCX erstellt: {OUTPUT_DOCX}")

    chapters = len(re.findall(r'^# Kapitel ', content, re.MULTILINE))
    print(f"Kapitel: {chapters}")
    print(f"Zeichen: {len(content):,}")

    print("Konvertiere zu PDF...")
    convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)


if __name__ == "__main__":
    main()
