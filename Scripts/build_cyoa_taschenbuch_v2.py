"""
Build a properly formatted DOCX for KDP Taschenbuch (paperback, 5x8 inches)
from the CYOA v2 (Time Cave) source files.

Reads graph_v2.yaml + all section .md files + frontmatter,
assembles them in randomized order (non-sequential to prevent skimming),
and generates a print-ready DOCX.

v2 changes vs. v1:
- String-based section IDs (P1, A1, B1, etc.) instead of integers
- Time Cave structure with path prefixes
- Sections stored in subdirectories (prolog/, pfad_a/, pfad_b/, etc.)
"""

import os
import re
import random
import yaml
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Benjamin Krug"
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.join(_SCRIPT_DIR, "..")
V2_DIR = os.path.join(_PROJECT_ROOT, "Band1", "CYOA", "v2")
GRAPH_FILE = os.path.join(V2_DIR, "graph_v2.yaml")
FRONTMATTER_FILE = os.path.join(_PROJECT_ROOT, "Band1", "CYOA", "frontmatter.md")
OUTPUT_FILE = os.path.join(_PROJECT_ROOT, "Output", "Band1", "KDP_Band1_CYOA_v2_Manuskript.docx")

SCENE_BREAK_SYMBOL = "\u2726  \u2726  \u2726"

# Seed for reproducible randomization
RANDOM_SEED = 42


def setup_page(section):
    """Set 5x8 inch page size and KDP margins."""
    section.page_width = Inches(5)
    section.page_height = Inches(8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.625)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.3)


def setup_mirror_margins(doc):
    """Enable mirror margins for bookbinding."""
    settings_element = doc.settings.element
    mirror = OxmlElement('w:mirrorMargins')
    settings_element.append(mirror)


def add_page_number_footer(section):
    """Add centered page number to section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run()
    run1.font.name = "Georgia"
    run1.font.size = Pt(10)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fld_begin)

    run2 = p.add_run()
    run2.font.name = "Georgia"
    run2.font.size = Pt(10)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._r.append(instr)

    run3 = p.add_run()
    run3.font.name = "Georgia"
    run3.font.size = Pt(10)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld_end)


def create_styles(doc):
    """Set up document styles for Taschenbuch formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.5)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    """Add a paragraph with inline formatting (*italic*, **bold**)."""
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)

    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.findall(pattern, text)

    for full_match, bold_text, italic_text, plain_text in parts:
        if bold_text:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif italic_text:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif plain_text:
            run = p.add_run(plain_text)
            run.font.name = "Georgia"
            run.font.size = Pt(11)

    return p


def add_section_heading(doc, print_num, is_ending=False, ending_title=None):
    """Add an Abschnitt heading with page break and spacing."""
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.keep_with_next = True

    if is_ending:
        title = ending_title or "Ende"
    else:
        title = "Abschnitt {}".format(print_num)

    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.name = "Georgia"
    run.bold = True


def add_front_matter(doc):
    """Add title page, half-title, copyright, and CYOA instructions."""
    add_blank_lines(doc, 6)
    add_centered_text(doc, "Die Geisterspuerer", font_size=18, bold=True)
    add_page_break(doc)

    add_blank_lines(doc, 3)
    add_centered_text(doc, "Die Geisterspuerer", font_size=22, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Das Haus, das fluestert", font_size=16, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Entscheidungsbuch", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Band 1", font_size=12)
    add_blank_lines(doc, 2)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    add_blank_lines(doc, 8)
    add_centered_text(
        doc,
        "Die Geisterspuerer - Das Haus, das fluestert (Entscheidungsbuch)",
        font_size=10,
        bold=True,
    )
    add_centered_text(doc, "Band 1", font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "© 2026 {}".format(AUTHOR), font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)


def add_cyoa_instructions(doc):
    """Add the 'how to read this book' page from frontmatter.md."""
    add_page_break(doc)

    if not os.path.exists(FRONTMATTER_FILE):
        print("WARNUNG: {} nicht gefunden, ueberspringe.".format(FRONTMATTER_FILE))
        return

    with open(FRONTMATTER_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    is_first = True
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            add_blank_lines(doc, 2)
            add_centered_text(doc, line[2:], font_size=16, bold=True)
            add_blank_lines(doc, 1)
            is_first = True
        elif line.startswith("## "):
            add_blank_lines(doc, 1)
            add_centered_text(doc, line[3:], font_size=13, bold=True)
            add_blank_lines(doc, 1)
            is_first = True
        elif line.startswith("---"):
            continue
        else:
            add_formatted_paragraph(doc, line, first_line_indent=not is_first)
            is_first = False


def add_back_matter(doc):
    """Add review request and series list."""
    add_page_break(doc)

    add_blank_lines(doc, 3)
    add_centered_text(
        doc,
        "\u201eDas Haus, das fluestert\u201c hat dir gefallen?",
        font_size=14,
        bold=True,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig ueber eine kurze Bewertung auf Amazon "
        "- auch nur ein oder zwei Saetze reichen voellig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Baende zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(
        doc, "Die Geisterspuerer \u2014 Alle Baende", font_size=14, bold=True
    )
    add_blank_lines(doc, 1)

    for band_num, title in [
        (1, "Das Haus, das fluestert"),
        (1, "Das Haus, das fluestert - Entscheidungsbuch"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tuer"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(
            doc,
            "**Band {}:** {}".format(band_num, title),
            first_line_indent=False,
        )


def load_graph():
    with open(GRAPH_FILE, "r", encoding="utf-8") as f:
        graph = yaml.safe_load(f)
    return graph


def load_section_content(file_path):
    """Load a markdown section file and return the body text (without heading)."""
    full_path = os.path.join(V2_DIR, file_path)
    if not os.path.exists(full_path):
        return None

    with open(full_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Remove the "# Abschnitt/Ende/Dead End" heading
    content = re.sub(r'^#[^\n]*\n+', '', content)

    # Remove subheadings
    content = re.sub(r'^##[^\n]*\n+', '', content, flags=re.MULTILINE)

    return content.strip()


def remap_references(content, sid_to_print_num):
    """Replace internal section references with print numbers.

    Transforms:
    - 'Abschnitt A1' -> 'Abschnitt 42'
    - 'Ende E7' -> 'Ende 87'
    - 'geh zu Abschnitt B16' -> 'geh zu Abschnitt 65'
    """
    def replace_ref(match):
        prefix = match.group(1)  # 'Abschnitt' or 'Ende'
        sid = match.group(2)     # e.g. 'A1', 'E7', 'B16'
        if sid in sid_to_print_num:
            return "{} {}".format(prefix, sid_to_print_num[sid])
        return match.group(0)  # Keep original if not found

    # Match patterns like "Abschnitt A1", "Ende E7", etc.
    content = re.sub(
        r'(Abschnitt|Ende)\s+([A-Z][A-Za-z_]*\d+[a-z]?)',
        replace_ref,
        content
    )

    return content


def randomize_section_order(section_ids, start_id="P1"):
    """Randomize section order for print, keeping P1 first."""
    remaining = [s for s in section_ids if s != start_id]
    random.seed(RANDOM_SEED)
    random.shuffle(remaining)
    return [start_id] + remaining


def build_section(doc, print_num, section_data, content):
    """Build a single CYOA section in the DOCX."""
    is_ending = section_data.get("type") == "ending"
    ending_title = section_data.get("title", "")

    if is_ending:
        add_section_heading(doc, print_num, is_ending=True, ending_title=ending_title)
    else:
        add_section_heading(doc, print_num)

    if not content:
        add_formatted_paragraph(doc, "[Inhalt fehlt]", first_line_indent=False)
        return

    lines = content.split("\n")
    is_first = True

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line == "---":
            add_blank_lines(doc, 1)
            add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
            add_blank_lines(doc, 1)
            is_first = True
            continue

        if line.startswith("**ENDE**"):
            add_blank_lines(doc, 1)
            add_centered_text(doc, "ENDE", font_size=14, bold=True)
            add_blank_lines(doc, 1)
            is_first = True
            continue

        add_formatted_paragraph(doc, line, first_line_indent=not is_first)
        is_first = False


def main():
    print("Lade graph_v2.yaml...")
    graph = load_graph()
    sections = graph["sections"]

    all_ids = sorted(sections.keys())
    print("Gefunden: {} Abschnitte in graph_v2.yaml".format(len(all_ids)))

    # Separate by type
    story_ids = [sid for sid in all_ids if sections[sid].get("type") != "ending"]
    ending_ids = [sid for sid in all_ids if sections[sid].get("type") == "ending"]

    print("  Story/Choice/Dead-End: {}".format(len(story_ids)))
    print("  Enden: {}".format(len(ending_ids)))

    # Randomize order
    print("Randomisiere Abschnitt-Reihenfolge...")
    ordered_story = randomize_section_order(story_ids)
    ordered_all = ordered_story + ending_ids

    # Create mapping: original ID -> print number
    sid_to_print_num = {}
    for i, sid in enumerate(ordered_all, start=1):
        sid_to_print_num[sid] = i

    # Load all content
    print("Lade Abschnitt-Inhalte...")
    contents = {}
    missing = []
    for sid in all_ids:
        sec = sections[sid]
        file_path = sec.get("file")
        if file_path:
            content = load_section_content(file_path)
            if content is None:
                missing.append(sid)
            else:
                # Remap internal references to print numbers
                content = remap_references(content, sid_to_print_num)
            contents[sid] = content
        else:
            missing.append(sid)
            contents[sid] = None

    if missing:
        print("WARNUNG: {} Abschnitte ohne Datei: {}".format(
            len(missing), missing[:20]))

    # Build DOCX
    print("Erstelle DOCX...")
    doc = Document()

    section_front = doc.sections[0]
    setup_page(section_front)

    header = section_front.header
    header.is_linked_to_previous = False
    footer = section_front.footer
    footer.is_linked_to_previous = False

    setup_mirror_margins(doc)
    create_styles(doc)
    add_front_matter(doc)
    add_cyoa_instructions(doc)

    # Body section with page numbers
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_body = doc.sections[-1]
    setup_page(section_body)
    add_page_number_footer(section_body)

    header_body = section_body.header
    header_body.is_linked_to_previous = False

    # Build all sections
    for sid in ordered_all:
        sec = sections[sid]
        content = contents.get(sid)
        print_num = sid_to_print_num[sid]
        build_section(doc, print_num, sec, content)

    add_back_matter(doc)

    doc.save(OUTPUT_FILE)
    print("\n{} erstellt!".format(OUTPUT_FILE))

    # Stats
    word_count = 0
    for content in contents.values():
        if content:
            word_count += len(content.split())

    print("Abschnitte: {}".format(len(all_ids)))
    print("Woerter (gesamt): ~{}".format(word_count))
    print("Fehlende Dateien: {}".format(len(missing)))

    # Print the mapping for reference
    print("\n--- Abschnitt-Zuordnung (intern -> Druck) ---")
    for sid in ordered_all[:10]:
        print("  {} -> Abschnitt {}".format(sid, sid_to_print_num[sid]))
    if len(ordered_all) > 10:
        print("  ... ({} weitere)".format(len(ordered_all) - 10))


if __name__ == "__main__":
    main()
