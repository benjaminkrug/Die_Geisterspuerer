"""
Build a properly formatted DOCX for KDP Taschenbuch (paperback, 6x9 inches)
from Manuskript_Band1_Komplett.md.
Handles: page size, mirror margins, page numbers, chapter headings,
paragraphs with first-line indent, scene breaks, italic/bold text,
front matter, back matter.

★ FORMATWECHSEL 2026-08-04: 5 x 8 Zoll -> 6 x 9 Zoll.

Band 1 lief als EINZIGER Band der Reihe auf 5 x 8 Zoll; Band 2 bis 5 sind
6 x 9 (build_taschenbuch_docx_band2..5.py). Im Regal stand der Einstiegsband
damit sichtbar kleiner als der Rest der Reihe, und das Cover konnte nie zum
Set passen.

Uebernommen wurden Seitengroesse, Raender und Erstzeileneinzug exakt aus
build_taschenbuch_docx_band5.py. Die uebrige Typografie war ohnehin identisch
(Georgia 11 pt, Zeilenabstand 1,5) — es aendert sich also nur das Format,
nicht das Schriftbild.

Folgen, die beachtet werden muessen:
  - Die SEITENZAHL sinkt (5 x 8 / 186 Seiten -> geschaetzt ~160). Davon haengt
    die Buchruecken-Breite ab: Scripts/build_cover.py braucht die echte Zahl
    aus dem KDP-Previewer, BEVOR ein Cover gebaut wird.
  - Das Taschenbuch muss bei KDP als NEUE Ausgabe hochgeladen werden; die ASIN
    unten (Rezensions-QR-Code) bitte gegenpruefen.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Benjamin Krug"
INPUT_FILE = os.path.join(os.path.dirname(__file__), "..", "Band1", "Manuskript", "Manuskript_Band1_Komplett.md")
OUTPUT_FILE = os.path.join(os.path.dirname(__file__), "..", "Output", "Band1", "KDP_Band1_Manuskript.docx")

SCENE_BREAK_SYMBOL = "\u2726  \u2726  \u2726"

# --- Amazon-Rezension: QR-Code auf der "Hat's dir gefallen?"-Seite ---
# ASIN von Band 1 (Taschenbuch). Bei neuer Ausgabe hier anpassen.
BAND1_ASIN = "B0GNZVXDDJ"
REVIEW_URL = f"https://www.amazon.de/review/create-review?asin={BAND1_ASIN}"
QR_IMAGE = os.path.join(os.path.dirname(__file__), "..", "Band1", "Cover", "qr_rezension_band1.png")


def setup_page(section):
    """Set 6x9 inch page size and KDP margins — identisch zu Band 2 bis 5."""
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = Inches(0.875)   # inner (binding)
    section.right_margin = Inches(0.625)  # outer
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.35)


def setup_mirror_margins(doc):
    """Enable mirror margins for bookbinding (inner/outer swap on even/odd pages)."""
    settings_element = doc.settings.element
    mirror = OxmlElement('w:mirrorMargins')
    settings_element.append(mirror)


def add_page_number_footer(section):
    """Add centered page number to section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PAGE field: begin
    run1 = p.add_run()
    run1.font.name = "Georgia"
    run1.font.size = Pt(10)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fld_begin)

    # PAGE field: instruction
    run2 = p.add_run()
    run2.font.name = "Georgia"
    run2.font.size = Pt(10)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._r.append(instr)

    # PAGE field: end
    run3 = p.add_run()
    run3.font.name = "Georgia"
    run3.font.size = Pt(10)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld_end)


def create_styles(doc):
    """Set up document styles for Taschenbuch formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.6)  # wie Band 2-5 (war 0.5 beim 5x8-Satz)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_lines(doc, count=1):
    """Add empty paragraphs."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_centered_image(doc, image_path, width_inches=1.6):
    """Add a centered image (e.g. QR code). Silently skips if file is missing."""
    if not os.path.isfile(image_path):
        print(f"  WARNUNG: Bild nicht gefunden, wird übersprungen: {image_path}")
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    """Add a paragraph with inline formatting (*italic*, **bold**)."""
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)

    # Parse inline markdown formatting
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.findall(pattern, text)

    for full_match, bold_text, italic_text, plain_text in parts:
        if bold_text:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif italic_text:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif plain_text:
            run = p.add_run(plain_text)
            run.font.name = "Georgia"
            run.font.size = Pt(11)

    return p


def add_front_matter(doc):
    """Add title page, half-title, and copyright page."""
    # --- Half title ---
    add_blank_lines(doc, 6)
    add_centered_text(doc, "Die Geisterspürer", font_size=18, bold=True)
    add_page_break(doc)

    # --- Full title page ---
    add_blank_lines(doc, 3)
    add_centered_text(doc, "Die Geisterspürer", font_size=22, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Das Haus, das flüstert", font_size=16, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Band 1", font_size=12)
    add_blank_lines(doc, 2)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # --- Copyright page ---
    add_blank_lines(doc, 8)
    add_centered_text(doc, "Die Geisterspürer – Das Haus, das flüstert", font_size=10, bold=True)
    add_centered_text(doc, "Band 1", font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = Pt(9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)


def add_back_matter(doc):
    """Add back matter: review request, Band 2 teaser, series list, cross-references."""
    add_page_break(doc)

    # --- 1. Rezensions-Bitte ---
    add_blank_lines(doc, 3)
    add_centered_text(doc, "\u201eDas Haus, das fl\u00fcstert\u201c hat dir gefallen?", font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 2)

    # QR-Code direkt zur Bewertung
    add_centered_text(doc, "Einfach den Code scannen und eine Bewertung dalassen:", font_size=11, italic=True)
    add_blank_lines(doc, 1)
    add_centered_image(doc, QR_IMAGE, width_inches=1.6)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "(Handykamera auf den Code halten – der Link öffnet sich von selbst.)", font_size=9, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # --- 2. Teaser Band 2 ---
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Weiterlesen? Hier kommt eine Vorschau auf Band 2:", font_size=11, italic=True)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Geisterspürer", font_size=18, bold=True)
    add_centered_text(doc, "Der Friedhof ohne Namen", font_size=14, italic=True)
    add_centered_text(doc, "Band 2", font_size=11)
    add_blank_lines(doc, 2)

    Q = "\u201e"
    E = "\u201c"
    teaser_paragraphs = [
        "Schatten wollte nicht rein.",
        "Das erfuhr Nora erst am Tor — aber beim Frühstück, eine Stunde vorher, hatte sie ihn schon angesehen und gedacht: Er weiß es schon. Er weiß immer, was kommt.",
        "Frau Silbers Karte lag seit Tagen auf ihrem Schreibtisch. Handgezeichnet, zwölf Markierungen. Eine war durchgestrichen. Lina. Elf blieben übrig.",
        "Die nächste: der alte Friedhof am Rand der Altstadt. Drei Straßen von ihrer Wohnung entfernt.",
        "Nora hatte beim Frühstück keine besondere Eile gezeigt.",
        "\"Du isst seit zwanzig Minuten dasselbe Brötchen\", sagte Theo.",
        "Er saß ihr gegenüber und hatte bereits zwei Brötchen, eine Schüssel Müsli und ein Glas Orangensaft geschafft. Schatten saß neben seinem Stuhl und wartete geduldig auf Krümel. Seine bernsteinfarbenen Augen gingen von Theo zu Nora und wieder zurück.",
        "\"Ich kaue gründlich\", sagte Nora.",
        "\"Du kaust gar nicht. Du starrst das Brötchen an.\"",
        "\"Ich denke nach.\"",
        "Theo lehnte sich zurück. \"Du hast Angst.\"",
        "Nora legte das Brötchen auf den Teller. Ihr Magen hatte sich beim Aufwachen schon zusammengezogen — ein kleines, unangenehmes Knoten-Gefühl, das sie die ganze Zeit beim Zähneputzen begleitet hatte. Ihr Mund war trocken. Den Tee hatte sie kaum angerührt.",
        "\"Ich habe keine Angst\", sagte sie.",
        "\"Wir gehen heute auf einen Friedhof\", sagte Theo. \"Um einen Geist zu suchen. Nach allem, was letzten Monat passiert ist.\" Er machte eine kleine Pause. \"Ich habe Angst. Das sage ich einfach mal laut.\"",
        "\"Dann bleib hier.\"",
        "\"Auf keinen Fall.\"",
        "Sie brachen zwanzig Minuten später auf.",
        "Die Sonne schien. Das war fast verdächtig. Nora hatte irgendwie mit Regen gerechnet, mit grauem Himmel, mit dem richtigen Wetter für einen Friedhofsbesuch. Stattdessen: klarer Herbstmorgen, goldenes Licht auf den Kopfsteinpflastern, der Geruch von frischen Brezeln aus der Bäckerei an der Ecke.",
        "Gravenstedt wirkte harmlos, wenn die Sonne schien.",
        "Nora wusste inzwischen, dass das täuschte.",
    ]

    for i, para in enumerate(teaser_paragraphs):
        add_formatted_paragraph(doc, para, first_line_indent=(i > 0))

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Jetzt überall erhältlich.", font_size=11, italic=True)

    # --- 3. Serien-Übersicht ---
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Geisterspürer \u2014 Alle Bände", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    for band_num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(
            doc,
            f"**Band {band_num}:** {title}",
            first_line_indent=False,
        )

    # Hinweis: Geisterspürer auch als interaktives Mitmach-Buch
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Übrigens: Die Geisterspürer gibt es auch als interaktives Mitmach-Buch, "
        "in dem *du* entscheidest, was Nora und Theo als Nächstes tun — mit vielen "
        "Wegen und Enden. Frag mal danach!",
        first_line_indent=False,
    )

    # --- 4. Cross-Verweis andere Serien ---
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Mehr spannende Abenteuer von Benjamin Krug:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    # Herrenhaus-Detektive
    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schl\u00fcssel \u2014 und hinter der verschlossenen "
        "T\u00fcr wartet ein Geheimnis, das seit drei\u00dfig Jahren niemand "
        "l\u00fcften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)

    # Herrenhaus interaktiv: "Dein Fall \u2013 Du entscheidest!"
    add_centered_text(doc, "Dein Fall \u2013 Du entscheidest!", font_size=13, bold=True)
    add_centered_text(doc, "Das verbotene Herrenhaus", font_size=11, italic=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Kein normales Buch: Hier bestimmst *du*, was passiert. An jeder wichtigen "
        "Stelle triffst du eine Entscheidung \u2013 und schl\u00e4gst die Seite auf, die zu "
        "deiner Wahl geh\u00f6rt.",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "Vier Wege durch das alte Herrenhaus, 14 verschiedene Enden. Findest du "
        "den richtigen? Oder tappst du in die Falle? Bei jedem Lesen wird die "
        "Geschichte neu.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Interaktives Detektiv-Spielbuch ab 8 Jahren.", font_size=11, italic=True)


def add_chapter_heading(doc, title):
    """Add a chapter heading with page break and precise spacing."""
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.name = "Georgia"
    run.bold = True


def add_scene_break(doc):
    """Add an ornamental scene break."""
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)


def parse_and_build(doc, content):
    """Parse markdown content and build the DOCX."""
    match = re.search(r'^# Kapitel 1', content, re.MULTILINE)
    if not match:
        raise ValueError("Could not find '# Kapitel 1' in manuscript")
    body = content[match.start():]

    # Remove ENDE BAND 1 and trailing ---
    body = re.sub(r'\n---\n\n\*\*ENDE BAND 1\*\*\n\n---\n*', '', body)
    body = re.sub(r'\n\*\*ENDE BAND 1\*\*\n*', '', body)

    lines = body.split('\n')

    def _naechste_zeile_ist_kapitel(start):
        """SEPARATOR-BUG-FIX (2026-07-18): schaut ueber Leerzeilen hinweg,
        ob als Naechstes eine Kapitelueberschrift kommt."""
        j = start
        while j < len(lines) and not lines[j].strip():
            j += 1
        return j < len(lines) and lines[j].strip().startswith('# Kapitel ')


    is_first_para_after_heading = False
    is_first_para_after_break = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Chapter heading
        if line.startswith('# Kapitel '):
            title = line[2:]  # Remove "# "
            add_chapter_heading(doc, title)
            is_first_para_after_heading = True
            i += 1
            continue

        # Scene break
        if line == '---':
            # ── SEPARATOR-BUG-FIX (2026-07-18) ────────────────────────────
            # Die Kompilier-Skripte setzen ein "---" ZWISCHEN die Kapitel.
            # Fuer diesen Parser ist das nicht von einem Szenentrenner zu
            # unterscheiden -> im Druck stand ein Ornament hinter JEDEM
            # Cliffhanger, dann erst der Seitenumbruch (Band 1: 17, B2: 14,
            # B3: 15, B4: 15 Stueck).
            # Ein Szenentrenner direkt vor einer Kapitelueberschrift ist
            # IMMER falsch -> ueberspringen.
            if _naechste_zeile_ist_kapitel(i + 1):
                i += 1
                continue
            add_scene_break(doc)
            is_first_para_after_break = True
            i += 1
            continue

        # Regular paragraph
        no_indent = is_first_para_after_heading or is_first_para_after_break
        add_formatted_paragraph(doc, line, first_line_indent=not no_indent)
        is_first_para_after_heading = False
        is_first_para_after_break = False
        i += 1


def main():
    # Read manuscript
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # --- Section 1: Front matter (no page numbers) ---
    section_front = doc.sections[0]
    setup_page(section_front)

    # Suppress footer on front matter
    header = section_front.header
    header.is_linked_to_previous = False
    footer = section_front.footer
    footer.is_linked_to_previous = False

    # Enable mirror margins for bookbinding
    setup_mirror_margins(doc)

    # Set up styles
    create_styles(doc)

    # Build front matter
    add_front_matter(doc)

    # --- Section 2: Body + back matter (with page numbers) ---
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_body = doc.sections[-1]
    setup_page(section_body)

    # Add page numbers to body section footer
    add_page_number_footer(section_body)

    # Suppress header in body section
    header_body = section_body.header
    header_body.is_linked_to_previous = False

    # Parse and build body content
    parse_and_build(doc, content)

    # Add back matter
    add_back_matter(doc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"{OUTPUT_FILE} erstellt!")

    # Stats
    para_count = len(doc.paragraphs)
    print(f"Absätze: {para_count}")

    chapters = len(re.findall(r'# Kapitel \d+', content))
    print(f"Kapitel: {chapters}")


if __name__ == "__main__":
    main()
