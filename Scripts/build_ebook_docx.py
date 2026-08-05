"""
Build a properly formatted DOCX for KDP eBook upload from Manuskript_Band1_Komplett.md.
Handles: chapter headings, paragraphs with first-line indent, scene breaks,
italic/bold text, front matter, back matter.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTHOR = "Benjamin Krug"
INPUT_FILE = os.path.join(os.path.dirname(__file__), "..", "Band1", "Manuskript", "Manuskript_Band1_Komplett.md")
OUTPUT_FILE = os.path.join(os.path.dirname(__file__), "..", "Output", "Band1", "KDP_Band1_eBook.docx")


def create_styles(doc):
    """Set up document styles for eBook formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.7)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Configure "Heading 1" so chapter titles are recognized by Kindle/KDP
    # for the auto-generated navigation TOC, while still looking like our
    # centered chapter headings. The key part is that chapters USE this style.
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Georgia"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = None  # inherit (black) instead of the default blue
    h1pf = h1.paragraph_format
    h1pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1pf.first_line_indent = Cm(0)
    h1pf.space_before = Pt(0)
    h1pf.space_after = Pt(0)
    h1pf.keep_with_next = True


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_blank_lines(doc, count=1):
    """Add empty paragraphs."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def add_formatted_paragraph(doc, text, first_line_indent=True, alignment=None):
    """Add a paragraph with inline formatting (*italic*, **bold**)."""
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if alignment:
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)

    # Parse inline markdown formatting
    # Handle **bold** and *italic*
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.findall(pattern, text)

    for full_match, bold_text, italic_text, plain_text in parts:
        if bold_text:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif italic_text:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(11)
        elif plain_text:
            run = p.add_run(plain_text)
            run.font.name = "Georgia"
            run.font.size = Pt(11)

    return p


def add_table_of_contents(doc):
    """Insert a real Word TOC field on its own page.

    KDP/Kindle reads this field (built from the Heading 1 chapter titles) and
    turns it into a clickable in-book table of contents. The field needs to be
    updated once in Word/LibreOffice (or by KDP on conversion) to fill in the
    page numbers, but the entries and links are generated automatically.
    """
    add_blank_lines(doc, 3)
    add_centered_text(doc, "Inhalt", font_size=18, bold=True)
    add_blank_lines(doc, 2)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()

    # Build the TOC field: \o "1-1" = use Heading 1 only,
    # \h = hyperlinks, \z = hide tab leader/page no in web view, \u = outline level
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # Placeholder text shown until the field is updated
    placeholder = OxmlElement("w:t")
    placeholder.text = "Inhaltsverzeichnis wird beim Öffnen aktualisiert."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_elem = run._r
    r_elem.append(fld_begin)
    r_elem.append(instr)
    r_elem.append(fld_sep)
    r_elem.append(placeholder)
    r_elem.append(fld_end)

    # Tell Word to update fields automatically when the document is opened,
    # so the TOC and its links are populated without manual interaction.
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    add_page_break(doc)


def add_front_matter(doc):
    """Add title page, half-title, and copyright page."""
    # --- Half title ---
    add_blank_lines(doc, 8)
    add_centered_text(doc, "Die Geisterspürer", font_size=18, bold=True)
    add_page_break(doc)

    # --- Full title page ---
    add_blank_lines(doc, 5)
    add_centered_text(doc, "Die Geisterspürer", font_size=22, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Das Haus, das flüstert", font_size=16, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Band 1", font_size=12)
    add_blank_lines(doc, 3)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # --- Copyright page ---
    add_blank_lines(doc, 12)
    add_centered_text(doc, "Die Geisterspürer – Das Haus, das flüstert", font_size=10, bold=True)
    add_centered_text(doc, "Band 1", font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = Pt(9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)
    add_page_break(doc)

    # --- Table of contents ---
    add_table_of_contents(doc)


def add_back_matter(doc):
    """Add back matter: review request, Band 2 teaser, series list, cross-references."""
    add_page_break(doc)

    # --- 1. Rezensions-Bitte ---
    add_blank_lines(doc, 3)
    add_centered_text(doc, 'Hat dir \u201eDas Haus, das fl\u00fcstert\u201c gefallen?', font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # --- 2. Teaser Band 2 ---
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Weiterlesen? Hier kommt eine Vorschau auf Band 2:", font_size=11, italic=True)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Geisterspürer", font_size=18, bold=True)
    add_centered_text(doc, "Der Friedhof ohne Namen", font_size=14, italic=True)
    add_centered_text(doc, "Band 2", font_size=11)
    add_blank_lines(doc, 2)

    # Teaser text — written in the same style as Band 1
    # German quotes: \u201e = opening „  \u201c = closing "
    Q = "\u201e"  # „
    E = "\u201c"  # "
    teaser_paragraphs = [
        # --- Sofort Spannung: Schatten reagiert ANDERS als sonst ---
        "Schatten wollte nicht auf den Friedhof.",
        f"Das war neu. Bei Lina hatte er sie nach oben gezerrt, in die verlassene Wohnung. Er hatte gekratzt und geknurrt und nicht aufgeh\u00f6rt, bis Nora ihm gefolgt war. Aber vor dem rostigen Eisentor des alten Friedhofs stand er still. Steif. Sein Fell ges\u00e4ubt, sein Blick starr \u2014 und zum ersten Mal wollte er zur\u00fcck.",
        f"{Q}Er hat Angst{E}, fl\u00fcsterte Theo.",
        f"Nora schluckte. Sie hatte Schatten knurren gesehen. Bellen. Sie hatte ihn vor sich und Theo stellen sehen wie einen Schild. Aber Angst? Noch nie.",
        # --- Emotionaler Haken: warum sie trotzdem reingehen ---
        f"{Q}Wir k\u00f6nnen umdrehen{E}, sagte Theo. Er meinte es ernst. Kein Witz diesmal.",
        f"Nora schaute auf die Karte in ihrer Hand. Frau Silbers Handschrift. *Heinrich M. \u2014 1847. W\u00fctend. Grab.* Hundertachtzig Jahre. Allein zwischen den Grabsteinen. W\u00e4hrend die Stadt um ihn herum lebte und lachte und ihn vergessen hatte.",
        f"{Q}Nein{E}, sagte sie. {Q}Wir gehen rein.{E}",
        f"{Q}Das sagt das M\u00e4dchen, das vor drei Wochen nicht mal an Geister geglaubt hat.{E}",
        f"{Q}Vor drei Wochen wusste ich auch nicht, dass unser Hund \u00fcbernat\u00fcrlich ist.{E}",
        # --- Sie betreten den Friedhof: kurz, schnell, atmosphärisch ---
        "Das Tor quietschte. Der Klang hallte zwischen den Grabsteinen.",
        "Nora ging voraus. Kies unter ihren Schuhen. Moos an den Steinen. Stille, die sich anf\u00fchlte wie Watte in den Ohren.",
        "Schatten folgte. Langsam. Jeder Schritt ein Kampf gegen seinen Instinkt.",
        # --- Mystery-Element: der namenlose Grabstein ---
        f"Dann blieb er stehen. Am Ende der dritten Reihe. Vor einem Grabstein, der anders war als alle anderen. Gr\u00f6\u00dfer. Dunkler. Schwarzer Stein, glatt poliert \u2014 als w\u00fcrde ihn jemand pflegen.",
        "Aber es stand kein Name darauf. Kein Datum. Nichts.",
        f"{Q}Ein Grabstein ohne Namen{E}, murmelte Theo. {Q}Wer liegt da?{E}",
        # --- Eskalation: etwas stimmt hier nicht ---
        "Nora kniete sich hin. Vor dem Stein lag etwas im Gras. Klein. Metallisch. Ein Schl\u00fcssel. Alt, verrostet, aber sorgf\u00e4ltig hingelegt. Als h\u00e4tte ihn jemand erst gestern dort platziert.",
        f"Und das Gras vor dem Grabstein war niedergedr\u00fcckt. In einem Kreis. Genau so gro\u00df wie ein Mensch, der steht.",
        f"{Q}Jemand steht hier{E}, fl\u00fcsterte sie. {Q}Jede Nacht. An genau dieser Stelle.{E}",
        f"{Q}Seit hundertachtzig Jahren{E}, sagte Theo. Seine Stimme war kaum zu h\u00f6ren.",
        # --- Schatten-Reaktion: der Hund weiß mehr ---
        f"Schatten hob den Kopf. Sein Knurren kam so tief aus seiner Brust, dass Nora es in den F\u00fc\u00dfen sp\u00fcrte. Sein ganzer K\u00f6rper vibrierte.",
        f"Dann tat er etwas, das er noch nie getan hatte. Er wich zur\u00fcck. Schritt f\u00fcr Schritt. Ohne den Blick von dem Grabstein zu nehmen.",
        # --- Cliffhanger: der Geist ist DA ---
        f"Die Temperatur fiel. Noras Atem wurde sichtbar. Wei\u00dfe W\u00f6lkchen vor ihrem Mund \u2014 mitten im Juli.",
        f"Und auf dem schwarzen Grabstein erschienen Buchstaben. Langsam. Einer nach dem anderen. Wie von einer unsichtbaren Hand in den Stein geritzt.",
        "G E H T.",
        f"Theo griff nach Noras Arm. Seine Finger waren eiskalt.",
        f"Die Buchstaben verschwanden. Neue kamen.",
        f"O D E R.",
        f"Stille. Schattens Knurren erstarb. Die Luft stand still. Sogar der Wind h\u00f6rte auf.",
        f"Dann das letzte Wort. Gr\u00f6\u00dfer als die anderen. Tiefer in den Stein gedr\u00fcckt.",
        "B L E I B T.",
        # --- Letzter Satz: Kaufimpuls ---
        f"Nora sp\u00fcrte ihr Herz bis in die Fingerspitzen. Hinter ihr wich Theo zur\u00fcck. Vor ihr leuchteten die Buchstaben im schwarzen Stein. Und neben ihr knurrte Schatten \u2014 nicht den Grabstein an.",
        "Sondern etwas hinter ihnen.",
    ]

    for i, para in enumerate(teaser_paragraphs):
        # First paragraph after heading: no indent
        add_formatted_paragraph(doc, para, first_line_indent=(i > 0))

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Erscheint 2026 auf Amazon.", font_size=11, italic=True)

    # --- 3. Serien-Übersicht ---
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Geisterspürer — Alle Bände", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    for band_num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(
            doc,
            f"**Band {band_num}:** {title}",
            first_line_indent=False,
        )

    # --- 4. Cross-Verweis andere Serien ---
    add_blank_lines(doc, 1)
    add_centered_text(doc, "* * *", font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Mehr spannende Abenteuer von Benjamin Krug:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    # Herrenhaus-Detektive
    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schl\u00fcssel \u2014 und hinter der verschlossenen "
        "T\u00fcr wartet ein Geheimnis, das seit drei\u00dfig Jahren niemand "
        "l\u00fcften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)

    # Chrono-Agenten
    add_centered_text(doc, "Die Chrono-Agenten", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Drei Kinder. Ein Tablet, das die Zeit \u00f6ffnet. Und ein "
        "Countdown, der nicht aufh\u00f6rt zu laufen.",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "Leo, Mila und Ben sind Agenten wider Willen \u2014 geschickt in die "
        "gef\u00e4hrlichsten Momente der Geschichte, um zu retten, was jemand "
        "zerst\u00f6ren will. Ihr Gegner hei\u00dft Codex. Und er kennt ihre Namen.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "F\u00fcr alle, die Geschichte spannend finden \u2014 und wissen wollen, "
        "was passiert, wenn man sie ver\u00e4ndert.",
        first_line_indent=False,
    )


def add_chapter_heading(doc, title):
    """Add a chapter heading with page break.

    Uses the real "Heading 1" style so Kindle/KDP picks the chapter up for the
    auto-generated navigation TOC and the in-book TOC field resolves to it.
    """
    add_page_break(doc)
    add_blank_lines(doc, 3)
    p = doc.add_paragraph(title, style="Heading 1")
    # Ensure the run font matches our book look (style sets this too, but be safe)
    for run in p.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(16)
        run.bold = True
    add_blank_lines(doc, 2)


def add_scene_break(doc):
    """Add a scene break (centered asterisks)."""
    add_blank_lines(doc, 1)
    add_centered_text(doc, "* * *", font_size=11)
    add_blank_lines(doc, 1)


def parse_and_build(doc, content):
    """Parse markdown content and build the DOCX."""
    # Skip the header (title, subtitle, word count, ---)
    match = re.search(r'^# Kapitel 1', content, re.MULTILINE)
    if not match:
        raise ValueError("Could not find '# Kapitel 1' in manuscript")
    body = content[match.start():]

    # Remove ENDE BAND 1 and trailing ---
    body = re.sub(r'\n---\n\n\*\*ENDE BAND 1\*\*\n\n---\n*', '', body)
    body = re.sub(r'\n\*\*ENDE BAND 1\*\*\n*', '', body)

    # Split into lines
    lines = body.split('\n')

    def _naechste_zeile_ist_kapitel(start):
        """SEPARATOR-BUG-FIX (2026-07-18): schaut ueber Leerzeilen hinweg,
        ob als Naechstes eine Kapitelueberschrift kommt."""
        j = start
        while j < len(lines) and not lines[j].strip():
            j += 1
        return j < len(lines) and lines[j].strip().startswith('# Kapitel ')


    is_first_para_after_heading = False
    is_first_para_after_break = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            # Empty line - skip
            i += 1
            continue

        # Chapter heading
        if line.startswith('# Kapitel '):
            title = line[2:]  # Remove "# "
            add_chapter_heading(doc, title)
            is_first_para_after_heading = True
            i += 1
            continue

        # Scene break
        if line == '---':
            # ── SEPARATOR-BUG-FIX (2026-07-18) ────────────────────────────
            # Die Kompilier-Skripte setzen ein "---" ZWISCHEN die Kapitel.
            # Fuer diesen Parser ist das nicht von einem Szenentrenner zu
            # unterscheiden -> im Druck stand ein Ornament hinter JEDEM
            # Cliffhanger, dann erst der Seitenumbruch (Band 1: 17, B2: 14,
            # B3: 15, B4: 15 Stueck).
            # Ein Szenentrenner direkt vor einer Kapitelueberschrift ist
            # IMMER falsch -> ueberspringen.
            if _naechste_zeile_ist_kapitel(i + 1):
                i += 1
                continue
            add_scene_break(doc)
            is_first_para_after_break = True
            i += 1
            continue

        # Regular paragraph
        no_indent = is_first_para_after_heading or is_first_para_after_break
        add_formatted_paragraph(doc, line, first_line_indent=not no_indent)
        is_first_para_after_heading = False
        is_first_para_after_break = False
        i += 1


def main():
    import docx.enum.text

    # Read manuscript
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up page margins (for eBook, use standard margins)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Set up styles
    create_styles(doc)

    # Build document
    add_front_matter(doc)
    parse_and_build(doc, content)
    add_back_matter(doc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"{OUTPUT_FILE} erstellt!")

    # Stats
    para_count = len(doc.paragraphs)
    print(f"Absätze: {para_count}")


if __name__ == "__main__":
    main()
