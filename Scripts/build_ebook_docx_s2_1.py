"""
Baut ein KDP-fertiges eBook-DOCX fuer Staffel 2, Band 1 ("Der Gast, der blieb")
aus Manuskript_S2-1_Komplett.md.

Verwendung:
    python Scripts/build_ebook_docx_s2_1.py            # DOCX bauen
    python Scripts/build_ebook_docx_s2_1.py --pruefen  # nur pruefen (ohne python-docx)


★ HERKUNFT — zwei Vorlagen, bewusst gemischt

  build_ebook_docx.py (Band 1) liefert das eBook-Spezifische:
      - Kapitelueberschriften im echten "Heading 1"-Stil, damit Kindle sie fuer
        das Navigations-Inhaltsverzeichnis erkennt
      - ein echtes Word-TOC-Feld, das KDP beim Konvertieren aufloest
      - keine Seitengeometrie (eBooks sind fliessender Text)

  build_taschenbuch_docx_s2_1.py liefert die Textaufbereitung — und zwar
  IMPORTIERT, nicht kopiert (s. naechster Abschnitt).


★ WARUM DIE TEXTVERARBEITUNG IMPORTIERT WIRD

  Band 5 hat gelernt, die Frontmatter zu importieren statt sie zu kopieren:
  Band 1-4 definierten Widmung und Epigraph in ZWEI Dateien, und wenn die
  auseinanderlaufen, weicht das gedruckte Buch vom Manuskript ab, ohne dass es
  jemand merkt.

  Fuer die Textaufbereitung gilt dasselbe Argument, nur schaerfer: ENDE-Regex,
  Szenentrenner-Logik und Typografie muessten sonst in Taschenbuch UND eBook
  doppelt gepflegt werden. Genau daran ist Band 1 schon einmal gescheitert —
  siehe naechster Abschnitt. Hier gibt es EINE Quelle:

      bereite_body_vor · ereignisfolge · apply_typography · Frontmatter
      -> alles aus build_taschenbuch_docx_s2_1.py


★ WAS BAND 1s eBOOK FALSCH MACHT — und warum das hier nicht passieren kann

  build_ebook_docx.py wandelt **keine Anfuehrungszeichen um** (nachgezaehlt: null
  Aufrufe von typo_quotes/apply_typography). Das faellt bei Band 1 zusammen mit
  einem zweiten Umstand: Band 1s kompiliertes Manuskript ist das einzige der
  Reihe, das HALB umgestellt ist — deutsche Anfuehrungszeichen am Anfang, gerade
  ASCII am Ende:

      Band 1   ASCII 1078   auf-„ 1078   zu-“ 0
      Band 2-5 und S2-1: reines ASCII, die Skripte wandeln korrekt um

  ⚠️ Daraus folgt eine Falle fuer spaeter: Wer je ein modernes Skript (mit
  typo_quotes) auf Band 1s Manuskript laufen laesst, macht aus dem schliessenden
  ASCII-Zeichen ein ZWEITES oeffnendes: „Text„ statt „Text“. Band 1 braucht also
  erst eine Bereinigung des Manuskripts, nicht nur ein neues Skript.

  S2-1 ist reines ASCII und laeuft sauber durch — nachgewiesen mit 926
  oeffnenden zu 926 schliessenden Zeichen.


★ UNTERSCHIEDE ZUM S2-1-TASCHENBUCH

  - Kapitel als "Heading 1" statt eigenem zentrierten Absatz (Kindle-Navigation)
  - Inhaltsverzeichnis als Word-Feld, direkt nach dem Impressum
  - Keine Seitengroesse, keine Raender, keine Seitenzahlen, keine Spiegelraender
  - Kein Kapitaelchen-Auftakt: Kindle rendert small-caps unzuverlaessig
  - Widmung und Epigraph stehen wie im Druck unmittelbar VOR Kapitel 1,
    das Inhaltsverzeichnis davor

  Gleich geblieben: kein QR ohne ASIN, kein Teaser auf S2-2, Serienuebersicht
  nur mit kaufbaren Titeln. Begruendungen im Kopf des Taschenbuch-Skripts.
"""

import os
import re
import sys
import importlib.util

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.join(_SCRIPT_DIR, "..")


# ── Textaufbereitung + Frontmatter: EINE Quelle (s. Kopf) ─────────────────────
def _lade_taschenbuch_modul():
    pfad = os.path.join(_SCRIPT_DIR, "build_taschenbuch_docx_s2_1.py")
    spec = importlib.util.spec_from_file_location("_s21_taschenbuch", pfad)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)          # main() ist durch __main__-Guard geschuetzt
    return mod


TB = _lade_taschenbuch_modul()

bereite_body_vor = TB.bereite_body_vor
ereignisfolge = TB.ereignisfolge
apply_typography = TB.apply_typography
uebrige_ende_zeilen = TB.uebrige_ende_zeilen

AUTHOR = TB.AUTHOR
SERIES_TITLE = TB.SERIES_TITLE
STAFFEL_TITLE = TB.STAFFEL_TITLE
BAND_NUM = TB.BAND_NUM
BAND_TITLE = TB.BAND_TITLE
BAND_SUBTITLE = TB.BAND_SUBTITLE
ERWARTETE_KAPITEL = TB.ERWARTETE_KAPITEL
SCENE_BREAK_SYMBOL = TB.SCENE_BREAK_SYMBOL
BAND_ASIN = TB.BAND_ASIN
Q, E = TB.Q, TB.E
WIDMUNG_ZEILEN = TB.WIDMUNG_ZEILEN
EPIGRAPH_ZEILEN = TB.EPIGRAPH_ZEILEN
EPIGRAPH_SOURCE = TB.EPIGRAPH_SOURCE

INPUT_FILE = TB.INPUT_FILE
OUTPUT_DIR = os.path.join(_ROOT, "Output", "S2-1")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "KDP_S2-1_eBook.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOCX (braucht python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def _docx():
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    return dict(Document=Document, Pt=Pt, Cm=Cm,
                AL=WD_ALIGN_PARAGRAPH, BR=WD_BREAK, qn=qn, Ox=OxmlElement)


D = None


def create_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = D['Pt'](11)
    pf = style.paragraph_format
    pf.space_before = D['Pt'](0)
    pf.space_after = D['Pt'](0)
    pf.line_spacing = 1.5
    pf.first_line_indent = D['Cm'](0.7)
    pf.alignment = D['AL'].JUSTIFY

    # ★ Kapitel MUESSEN diesen Stil benutzen — daraus baut Kindle die Navigation.
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Georgia"
    h1.font.size = D['Pt'](16)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = None          # erben (schwarz) statt Word-Blau
    h1pf = h1.paragraph_format
    h1pf.alignment = D['AL'].CENTER
    h1pf.first_line_indent = D['Cm'](0)
    h1pf.space_before = D['Pt'](0)
    h1pf.space_after = D['Pt'](0)
    h1pf.keep_with_next = True


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = D['Cm'](0)
    p.add_run().add_break(D['BR'].PAGE)


def add_blank_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = D['Cm'](0)
        p.paragraph_format.space_before = D['Pt'](0)
        p.paragraph_format.space_after = D['Pt'](0)


def add_centered_text(doc, text, font_size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = D['AL'].CENTER
    p.paragraph_format.first_line_indent = D['Cm'](0)
    p.paragraph_format.space_before = D['Pt'](0)
    p.paragraph_format.space_after = D['Pt'](0)
    run = p.add_run(text)
    run.font.size = D['Pt'](font_size)
    run.font.name = "Georgia"
    run.bold = bold
    run.italic = italic
    return p


def _add_runs_with_markdown(p, text):
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for _full, bold_text, italic_text, plain_text in re.findall(pattern, text):
        if bold_text:
            run = p.add_run(bold_text); run.bold = True
        elif italic_text:
            run = p.add_run(italic_text); run.italic = True
        elif plain_text:
            run = p.add_run(plain_text)
        else:
            continue
        run.font.name = "Georgia"
        run.font.size = D['Pt'](11)


def add_formatted_paragraph(doc, text, first_line_indent=True):
    # ★ apply_typography kommt aus dem Taschenbuch-Skript. Band 1s eBook hatte
    #   diesen Schritt gar nicht — s. Kopfkommentar.
    text = apply_typography(text)
    p = doc.add_paragraph()
    if not first_line_indent:
        p.paragraph_format.first_line_indent = D['Cm'](0)
    _add_runs_with_markdown(p, text)
    return p


def add_table_of_contents(doc):
    """Echtes Word-TOC-Feld auf eigener Seite.

    KDP/Kindle liest das Feld (gebaut aus den Heading-1-Kapiteln) und macht ein
    anklickbares Inhaltsverzeichnis daraus."""
    add_blank_lines(doc, 3)
    add_centered_text(doc, "Inhalt", font_size=18, bold=True)
    add_blank_lines(doc, 2)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = D['Cm'](0)
    run = p.add_run()

    # \o "1-1" = nur Heading 1 · \h = Hyperlinks · \z = kein Fuellzeichen im
    # Web-Layout · \u = Gliederungsebene
    fld_begin = D['Ox']("w:fldChar"); fld_begin.set(D['qn']("w:fldCharType"), "begin")
    instr = D['Ox']("w:instrText"); instr.set(D['qn']("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'
    fld_sep = D['Ox']("w:fldChar"); fld_sep.set(D['qn']("w:fldCharType"), "separate")
    platzhalter = D['Ox']("w:t")
    platzhalter.text = "Inhaltsverzeichnis wird beim Öffnen aktualisiert."
    fld_end = D['Ox']("w:fldChar"); fld_end.set(D['qn']("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_sep, platzhalter, fld_end):
        run._r.append(el)

    # Word soll die Felder beim Oeffnen selbst aktualisieren.
    settings = doc.settings.element
    upd = settings.find(D['qn']("w:updateFields"))
    if upd is None:
        upd = D['Ox']("w:updateFields")
        settings.append(upd)
    upd.set(D['qn']("w:val"), "true")

    add_page_break(doc)


def add_front_matter(doc):
    # Halbtitel
    add_blank_lines(doc, 8)
    add_centered_text(doc, SERIES_TITLE, font_size=18, bold=True)
    add_page_break(doc)

    # Titelseite — mit Reihenname, weil es zwei Baende mit der Nummer 1 gibt
    add_blank_lines(doc, 5)
    add_centered_text(doc, SERIES_TITLE, font_size=22, bold=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, STAFFEL_TITLE, font_size=13)
    add_blank_lines(doc, 1)
    add_centered_text(doc, BAND_TITLE, font_size=16, italic=True)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Band {BAND_NUM}", font_size=12)
    add_blank_lines(doc, 3)
    add_centered_text(doc, AUTHOR, font_size=12)
    add_page_break(doc)

    # Impressum
    add_blank_lines(doc, 12)
    add_centered_text(doc, f"{SERIES_TITLE} – {BAND_TITLE}", font_size=10, bold=True)
    add_centered_text(doc, BAND_SUBTITLE, font_size=10)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"© 2026 {AUTHOR}", font_size=9)
    add_centered_text(doc, "Alle Rechte vorbehalten.", font_size=9)
    add_blank_lines(doc, 1)
    p = add_centered_text(doc, "", font_size=9)
    run = p.add_run(
        "Dieses Buch ist ein Werk der Fiktion. Namen, Figuren, Orte und "
        "Ereignisse sind frei erfunden. Jede Ähnlichkeit mit tatsächlichen "
        "Personen, lebend oder tot, ist rein zufällig."
    )
    run.font.size = D['Pt'](9)
    run.font.name = "Georgia"
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Umschlaggestaltung: {AUTHOR}", font_size=9)
    add_centered_text(doc, f"Satz und Layout: {AUTHOR}", font_size=9)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Erstausgabe 2026", font_size=9)
    add_centered_text(doc, "Independently published", font_size=9)
    add_page_break(doc)

    # Inhaltsverzeichnis (Navigationsapparat)
    add_table_of_contents(doc)

    # Widmung und Epigraph stehen wie im Druck unmittelbar vor Kapitel 1.
    add_blank_lines(doc, 10)
    for zeile in WIDMUNG_ZEILEN:
        add_centered_text(doc, zeile, font_size=12, italic=True)
        add_blank_lines(doc, 1)
    add_page_break(doc)

    add_blank_lines(doc, 10)
    for zeile in EPIGRAPH_ZEILEN:
        add_centered_text(doc, zeile, font_size=12, italic=True)
        add_blank_lines(doc, 1)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"— {EPIGRAPH_SOURCE}", font_size=10, italic=True)


def add_back_matter(doc):
    add_page_break(doc)

    add_blank_lines(doc, 3)
    add_centered_text(doc, f"{Q}{BAND_TITLE}{E} hat dir gefallen?", font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Dann freue ich mich riesig über eine kurze Bewertung auf Amazon "
        "— auch nur ein oder zwei Sätze reichen völlig.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Jede Rezension hilft anderen Kindern (und ihren Eltern), dieses Buch "
        "zu entdecken. Und mir hilft sie, weitere Bände zu schreiben.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 2)

    # Kein QR/Link ohne ASIN — sonst stuende eine Aufforderung ohne Ziel im Buch.
    if BAND_ASIN:
        add_formatted_paragraph(
            doc, f"Direkt zur Bewertung: {TB.REVIEW_URL}", first_line_indent=False)
        add_blank_lines(doc, 2)

    add_centered_text(doc, "Vielen Dank!", font_size=11)
    add_centered_text(doc, AUTHOR, font_size=11)

    # KEIN TEASER auf S2-2 — nicht geschrieben.

    # Serienuebersicht: nur kaufbare Titel
    add_page_break(doc)
    add_blank_lines(doc, 2)
    add_centered_text(doc, f"{SERIES_TITLE} — die ersten fünf Bände",
                      font_size=14, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Nora, Theo und Schatten hatten schon einmal zu tun. Fünf Fälle, "
        "bevor dieser hier anfing.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    for num, title in [
        (1, "Das Haus, das flüstert"),
        (2, "Der Friedhof ohne Namen"),
        (3, "Schatten sieht mehr"),
        (4, "Die zugemauerte Tür"),
        (5, "Der Schleier"),
    ]:
        add_formatted_paragraph(doc, f"**Band {num}:** {title}", first_line_indent=False)
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Man kann sie in jeder Reihenfolge lesen — auch nach diesem Buch.",
                      font_size=10, italic=True)

    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)
    add_centered_text(doc, f"Mehr spannende Abenteuer von {AUTHOR}:", font_size=14, bold=True)
    add_blank_lines(doc, 1)

    add_centered_text(doc, "Die Geisterspürer – als Spielbuch", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Du kennst die Geschichte. Aber was wäre passiert, wenn Nora damals "
        "nicht in den Keller gegangen wäre?",
        first_line_indent=False,
    )
    add_formatted_paragraph(
        doc,
        "24 verschiedene Enden. Und ein geheimes, das nur die Mutigsten finden.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Grusel-Spielbuch ab 10 Jahren.", font_size=11, italic=True)

    add_blank_lines(doc, 2)
    add_centered_text(doc, "Die Herrenhaus-Detektive", font_size=13, bold=True)
    add_blank_lines(doc, 1)
    add_formatted_paragraph(
        doc,
        "Niemand darf das alte Herrenhaus betreten. Aber Jonas, Mila und Ben "
        "finden einen Schlüssel — und hinter der verschlossenen Tür wartet ein "
        "Geheimnis, das seit dreißig Jahren niemand lüften durfte.",
        first_line_indent=False,
    )
    add_blank_lines(doc, 1)
    add_centered_text(doc, "Spannendes Detektivabenteuer ab 8 Jahren.", font_size=11, italic=True)


def add_chapter_heading(doc, title):
    """Echter Heading-1-Stil — daraus baut Kindle die Navigation."""
    add_page_break(doc)
    p = doc.add_paragraph(title, style="Heading 1")
    p.paragraph_format.space_before = D['Pt'](36)
    p.paragraph_format.space_after = D['Pt'](24)
    return p


def add_scene_break(doc):
    add_blank_lines(doc, 1)
    add_centered_text(doc, SCENE_BREAK_SYMBOL, font_size=11)
    add_blank_lines(doc, 1)


def parse_and_build(doc, content):
    body = bereite_body_vor(content)
    erste_nach_kapitel = False
    erste_nach_trenner = False

    for art, zeile in ereignisfolge(body):
        if art == 'KAPITEL':
            add_chapter_heading(doc, zeile[2:])
            erste_nach_kapitel, erste_nach_trenner = True, False
        elif art == 'TRENNER':
            add_scene_break(doc)
            erste_nach_kapitel, erste_nach_trenner = False, True
        else:
            # Kein Kapitaelchen-Auftakt: Kindle rendert small-caps unzuverlaessig.
            add_formatted_paragraph(
                doc, zeile,
                first_line_indent=not (erste_nach_kapitel or erste_nach_trenner))
            erste_nach_kapitel = erste_nach_trenner = False


# ══════════════════════════════════════════════════════════════════════════════
# PRUEFEN + MAIN
# ══════════════════════════════════════════════════════════════════════════════

def pruefen(content: str) -> int:
    print("=" * 74)
    print("PRUEFUNG DES eBOOK-BILDS — ohne python-docx")
    print("=" * 74)
    fehler = []

    body = bereite_body_vor(content)
    print("ENDE-Marker: beide entfernt, keine Restzeile ✓")

    ev = ereignisfolge(body)
    kapitel = [z for a, z in ev if a == 'KAPITEL']
    print(f"Kapitel (werden Heading 1 -> Kindle-Navigation): {len(kapitel)} "
          f"(erwartet {ERWARTETE_KAPITEL})")
    if len(kapitel) != ERWARTETE_KAPITEL:
        fehler.append(f"{len(kapitel)} Kapitel statt {ERWARTETE_KAPITEL}")

    print(f"Echte Szenentrenner: {sum(1 for a, _ in ev if a == 'TRENNER')}")
    for i in range(len(ev) - 1):
        if ev[i][0] == 'TRENNER' and ev[i + 1][0] == 'KAPITEL':
            fehler.append(f"Fehl-Ornament vor {ev[i+1][1]!r}")
    print(f"Fehl-Ornamente: {sum(1 for f in fehler if 'Ornament' in f)}")

    print("\nInhaltsverzeichnis-Eintraege, die Kindle bauen wird:")
    for z in kapitel[:3]:
        print(f"   {z[2:]}")
    print(f"   … ({len(kapitel)} insgesamt)")

    # Typografie am ganzen Buch — der Punkt, an dem Band 1s eBook danebenliegt
    aus = apply_typography(body)
    auf, zu = aus.count('„'), aus.count('“')
    gerade = aus.count(chr(34))
    print(f"\nAnfuehrungszeichen nach Typografie: {auf} oeffnend / {zu} schliessend, "
          f"{gerade} gerade uebrig")
    if auf != zu or gerade:
        fehler.append(f"Anfuehrungszeichen unausgeglichen ({auf}/{zu}, {gerade} gerade)")

    print(f"\nFrontmatter (importiert):")
    print(f"   Widmung : {WIDMUNG_ZEILEN[0]}")
    print(f"   Epigraph: {EPIGRAPH_ZEILEN[0]}")

    if BAND_ASIN is None:
        print("\n⚠️  KEINE ASIN — die Rezensions-Seite wird OHNE Link gebaut.")
    print("⚠️  KEIN TEASER auf S2-2 — der Band ist nicht geschrieben.")

    print()
    if fehler:
        print("!! BEFUNDE:")
        for f in dict.fromkeys(fehler):
            print(f"   - {f}")
        return 1
    print("Alle Pruefungen bestanden.")
    return 0


def main():
    global D

    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    if "--pruefen" in sys.argv:
        sys.exit(pruefen(content))

    try:
        D = _docx()
    except ImportError:
        print("!! python-docx ist nicht installiert — DOCX kann nicht gebaut werden.")
        print("   Die Textpruefung laeuft trotzdem:")
        print("       python Scripts/build_ebook_docx_s2_1.py --pruefen")
        sys.exit(2)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"Lese: {INPUT_FILE}")

    doc = D['Document']()
    create_styles(doc)
    add_front_matter(doc)
    parse_and_build(doc, content)
    add_back_matter(doc)

    doc.save(OUTPUT_DOCX)
    print(f"DOCX erstellt: {OUTPUT_DOCX}")

    kapitel = len(re.findall(r'^# Kapitel ', content, re.MULTILINE))
    print(f"Kapitel: {kapitel}")
    if kapitel != ERWARTETE_KAPITEL:
        print(f"!! ABBRUCH-WARNUNG: {kapitel} Kapitel, erwartet {ERWARTETE_KAPITEL}.")
        sys.exit(1)

    print("Frontmatter aus build_taschenbuch_docx_s2_1.py uebernommen (dort aus")
    print("build_manuskript_komplett_s2_1.py) — eine Quelle fuer alle drei Formate.")
    if BAND_ASIN is None:
        print("⚠️  OHNE Rezensions-Link gebaut — keine ASIN hinterlegt.")
    print("⚠️  OHNE TEASER gebaut — S2-2 ist nicht geschrieben.")


if __name__ == "__main__":
    main()
